"""Build the frozen AXYS ML final project report (Word). Read-only of results."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips, Cm, Emu

NAVY = RGBColor(0x1B, 0x3A, 0x4B)
TEAL = RGBColor(0x2F, 0x5D, 0x62)
RULE = RGBColor(0x6B, 0x4C, 0x7A)
GRAY = RGBColor(0x44, 0x44, 0x44)
HEADER_SHADING = "1B3A4B"
ROW_ALT = "F4F1F5"
ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "AXYS_ML_Final_Project_Report.docx"


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), kwargs.get("val", "single"))
        el.set(qn("w:sz"), kwargs.get("sz", "4"))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), kwargs.get("color", "BFB3C7"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def set_table_widths(table, widths_in):
    table.autofit = False
    total = sum(widths_in)
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(total * 1440)))
    tblW.set(qn("w:type"), "dxa")
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for child in list(grid):
            grid.remove(child)
    else:
        grid = OxmlElement("w:tblGrid")
        tblPr.addnext(grid)
    for w in widths_in:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 1440)))
        grid.append(gc)
    for row in table.rows:
        for cell, w in zip(row.cells, widths_in):
            cell.width = Inches(w)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(int(w * 1440)))
            tcW.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    r = run._r
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "separate")
    hint = OxmlElement("w:t")
    hint.text = "Right-click and choose Update Field to refresh the table of contents."
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "end")
    r.append(fc1)
    r.append(instr)
    r.append(fc2)
    r.append(hint)
    r.append(fc3)


class Report:
    def __init__(self):
        self.doc = Document()
        self._setup()

    def _setup(self):
        sec = self.doc.sections[0]
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.font.color.rgb = GRAY
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.line_spacing = 1.15
        for name, size, before, after, color in (
            ("Title", 26, 0, 6, NAVY),
            ("Heading 1", 16, 18, 8, NAVY),
            ("Heading 2", 13, 14, 6, TEAL),
            ("Heading 3", 12, 10, 4, RULE),
        ):
            st = styles[name]
            st.font.name = "Calibri"
            st.font.size = Pt(size)
            st.font.bold = True
            st.font.color.rgb = color
            st.paragraph_format.space_before = Pt(before)
            st.paragraph_format.space_after = Pt(after)
            st.paragraph_format.keep_with_next = True
        footer = sec.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("AXYS ML  ·  Final Project Report  ·  Page ")
        set_run_font(r, size=9, color=RGBColor(0x77, 0x77, 0x77))
        add_page_number(p)
        r2 = p.add_run("  ·  Frozen analyses; numbers from primary reports")
        set_run_font(r2, size=9, color=RGBColor(0x77, 0x77, 0x77))

    def h(self, text, level=1):
        return self.doc.add_heading(text, level=level)

    def p(self, text, *, italic=False, bold=False, size=11, space_after=8, align=None):
        para = self.doc.add_paragraph()
        if align:
            para.alignment = align
        para.paragraph_format.space_after = Pt(space_after)
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=GRAY)
        return para

    def quote(self, text):
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.35)
        para.paragraph_format.right_indent = Inches(0.2)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(10)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), "6B4C7A")
        pBdr.append(left)
        pPr.append(pBdr)
        run = para.add_run(text)
        set_run_font(run, size=11, italic=True, color=NAVY)
        return para

    def bullets(self, items):
        for it in items:
            para = self.doc.add_paragraph(it, style="List Bullet")
            para.paragraph_format.space_after = Pt(3)
            for run in para.runs:
                set_run_font(run, size=11, color=GRAY)

    def table(self, headers, rows, widths=None, caption=None):
        if caption:
            cap = self.doc.add_paragraph()
            cap.paragraph_format.space_after = Pt(4)
            cap.paragraph_format.space_before = Pt(8)
            r = cap.add_run(caption)
            set_run_font(r, size=9, italic=True, color=TEAL)
        n = len(headers)
        if widths is None:
            widths = [6.5 / n] * n
        tbl = self.doc.add_table(rows=1 + len(rows), cols=n)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        set_table_widths(tbl, widths)
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(h)
            set_run_font(run, size=9, bold=True, color=RGBColor(255, 255, 255))
            shade_cell(cell, HEADER_SHADING)
            set_cell_border(cell, color="1B3A4B", sz="4")
        prevent_row_split(tbl.rows[0])
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = tbl.rows[ri + 1].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                set_run_font(run, size=9, color=GRAY)
                if ri % 2 == 1:
                    shade_cell(cell, ROW_ALT)
                set_cell_border(cell, color="D9D0E0", sz="4")
            prevent_row_split(tbl.rows[ri + 1])
        sp = self.doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(6)
        return tbl

    def h3_block(self, letter, title):
        self.h(f"{letter}. {title}", 3)

    def step(
        self,
        heading,
        question,
        inp,
        method,
        prereg,
        result_prose,
        table_caption,
        headers,
        rows,
        widths,
        decision,
        artifacts,
        extra_prose=None,
        extra_tables=None,
    ):
        self.h(heading, 2)
        self.h3_block("a", "Question")
        self.p(question)
        self.h3_block("b", "Input")
        self.p(inp)
        self.h3_block("c", "Method")
        for para in method:
            self.p(para)
        self.h3_block("d", "Pre-registration")
        self.p(prereg)
        self.h3_block("e", "Result")
        for para in result_prose:
            self.p(para)
        self.table(headers, rows, widths=widths, caption=table_caption)
        if extra_prose:
            for para in extra_prose:
                self.p(para)
        if extra_tables:
            for cap, hh, rr, ww in extra_tables:
                self.table(hh, rr, widths=ww, caption=cap)
        self.h3_block("f", "Decision")
        for para in decision:
            self.p(para)
        self.h3_block("g", "Output artifacts")
        self.bullets(artifacts)


def build():
    r = Report()
    d = r.doc

    # ----- Title page -----
    for _ in range(3):
        d.add_paragraph()
    t = d.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("AXYS ML")
    set_run_font(run, size=32, bold=True, color=NAVY)
    st = d.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = st.add_run("Final Project Report")
    set_run_font(run, size=22, bold=True, color=TEAL)
    sub = d.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Female gait kinematics and victimization history:\n"
        "subject-level signature discovery, shared-pattern tests, and power"
    )
    set_run_font(run, size=13, italic=True, color=GRAY)
    meta = d.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(24)
    run = meta.add_run(
        "Analyses frozen 13–17 August 2026\n"
        "n = 31 independent subjects  ·  17 victimized / 14 control\n"
        "Phases 0–6  ·  Similarity P0.1–P0.6  ·  P0.1 power / MDE\n"
        "P1 (Wasserstein / RV / soft-DTW / common subspace) not started"
    )
    set_run_font(run, size=11, color=GRAY)
    note = d.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(36)
    run = note.add_run(
        "This document is synthesis only. No analysis was recomputed, "
        "re-run, or reinterpreted. Every numeric claim is taken from the "
        "primary report, certification, or locked JSON/CSV cited in that section. "
        "Where a later summary (README or p0_synthesis.md) rounds a value, "
        "this report uses the primary file’s precision and notes the rounding."
    )
    set_run_font(run, size=10, italic=True, color=TEAL)
    d.add_page_break()

    r.h("Contents", 1)
    r.p(
        "Microsoft Word will populate the entries below when the TOC field is updated "
        "(References → Table of Contents → Update Table, or right-click the field)."
    )
    toc_p = d.add_paragraph()
    add_toc_field(toc_p)
    d.add_page_break()

    # ============================================================
    # 1. Executive summary
    # ============================================================
    r.h("1. Executive Summary", 1)
    r.p(
        "This report is the final written account of the AXYS ML gait research program. "
        "It does not introduce new analyses. It walks through every completed step — "
        "dataset audit through cycle extraction, features, group statistics, unsupervised "
        "phenotypes, within-victim Euclidean structure, time-resolved trajectories, the "
        "six-test Similarity P0 battery, and the post-hoc P0.1 power / minimum-detectable-effect "
        "(MDE) simulation — stating for each the question, lineage of inputs, method, "
        "locked choices, numbers, decision, and files a reader can open to verify."
    )
    r.h("The two research questions", 2)
    r.p("Primary question (Phases 0–6), stated exactly as in README §1:")
    r.quote(
        "Which biomechanical gait characteristics, if any, show statistically supported, "
        "robust, interpretable, and subject-consistent differences between victimized and "
        "non-victimized females?"
    )
    r.p("Secondary question (Similarity P0 — after the first layer was null), stated exactly as in README §1:")
    r.quote(
        "Do the 17 victimized women share a locomotor pattern with each other that controls "
        "do not — a shared deviation direction, abnormality set, waveform shape, "
        "phase-localized signature, or inter-joint coupling — that univariate mean "
        "differences can miss?"
    )
    r.p(
        "Neither question is the same as “Can a model predict victimization from gait at high "
        "accuracy?” With 31 people and thousands of features, a classifier can overfit easily. "
        "The repository first asks whether a shared pattern exists. If the data do not show one, "
        "that result is reported honestly."
    )

    r.h("Answer to the primary question (Phases 0–6)", 2)
    r.p(
        "No statistically supported, robust, interpretable, and subject-consistent "
        "victim-versus-control gait difference was detected in this sample. Phase 3 tested "
        "335 label-blind redundancy representatives with Mann–Whitney tests, Cliff’s δ, "
        "Benjamini–Hochberg FDR, leave-one-subject-out (LOSO) direction, and victim "
        "directional consistency: 0 features met FDR q ≤ 0.10, 0 met q ≤ 0.05, and 0 met "
        "the pre-specified signature rule. Phase 4 found a stable hierarchical split that is "
        "a 27-versus-4 majority/outgroup partition, not two large gait types, with victimization "
        "enrichment Fisher p = 1 and permutation p = 1. Phase 5 found that the 17 victims are "
        "not more compact in the certified 27-D family-PC space than random groups of 17 "
        "(mean pairwise distance 20.4168 versus null mean 20.0601, permutation p = 0.547) "
        "and that no stable victim subgroup existed. Phase 6 found 0 ROBUST and 0 EXPLORATORY "
        "time-resolved trajectory regions after subject-level cluster permutation. Certification "
        "status for these inferential phases is PASS WITH WARNINGS, the warning being the "
        "honest absence of a manufactured signature."
    )

    r.h("Answer to the secondary question (Similarity P0)", 2)
    r.p(
        "No defensible shared victim locomotor signature was detected under the pre-registered "
        "P0 program. After residualization on height, mass, mean leg length, and subject-median "
        "cycle duration, every discovery gate was null: P0.1 mean pairwise cosine 0.0518 versus "
        "null mean 0.1103, permutation p = 0.7579; P0.2 mean pairwise Jaccard 0.1906 versus "
        "null 0.2012, p = 0.6032, 0/30 co-exceedance FDR survivors; P0.3 Pearson −0.0234 "
        "(p = 0.2669) and DTW 7.6570 (p = 0.6288), 0/12 FDR survivors; P0.4 0/240 cells with "
        "FDR q ≤ 0.10 (minimum raw permutation p = 0.0188); P0.5 was not a separate discovery "
        "test (residualization was the primary gate inside the others); P0.6 circular similarity "
        "0.6494 (p = 0.929) and DTW 17.6251 (p = 0.6048), 0/12 FDR survivors. High absolute "
        "shape or CRP similarity, where it occurs, matches the subject-label permutation null "
        "and is generic gait, not victim-specific excess."
    )

    r.h("Headline power / MDE finding", 2)
    r.p(
        "A post-hoc P0.1 simulation on the frozen residualized 31 × 27 Phase 4 cloud "
        "(empirical injection; 1000 datasets per λ; 999 permutations per dataset; seed 20260813) "
        "had 80% power only at λ = 0.7285 (report headline λ = 0.73), that is, a shared "
        "deviation-direction magnitude of 0.73× the typical control residual deviation-vector "
        "norm (median control ||d_i|| = 8.3830). At that λ, simulated mean pairwise cosine is "
        "approximately 0.300. The observed residualized cosine 0.0518 is far below that "
        "threshold. At a more modest shared shift (λ = 0.50) simulated power was only 0.291. "
        "False-positive rate at λ = 0 was 0.052; power at λ = 3.0 was 1.000; the simulated "
        "λ = 0 cosine distribution matched the real residualized P0.1 permutation null "
        "(mean 0.109 versus 0.110). This design was powered for a gross shared direction, "
        "not a subtle one. Absence of evidence is not proof that no gait difference exists "
        "in a larger population."
    )

    r.h("What was not tested (scope)", 2)
    r.p(
        "This summary cannot be over-read as a clinical, causal, or predictive claim. "
        "The program did not: train a victim classifier, XGBoost, neural net, or “victim score”; "
        "treat 880 gait cycles as independent samples; start Similarity P1 (Wasserstein, RV, "
        "soft-DTW, common subspace); start Phase 7 supervised prediction; test males (removed "
        "from the processed file); capture daily walking (lab Plug-in Gait at 100 Hz); "
        "name coordinate axes as anterior–posterior / mediolateral / vertical; reconstruct "
        "initial, mid, or terminal swing for P0.4; estimate power for P0.2–P0.6; or treat "
        "survey victimization as a controlled exposure. Victim type among the 17 labeled Y is "
        "heterogeneous (Nd 7, in-person Ip 6, Both 3, online 1). Those untested constructs "
        "are limitations, not implied negatives beyond the tests that were actually run."
    )

    # ============================================================
    # 2. Background
    # ============================================================
    r.h("2. Background and Motivation", 1)
    r.h("Why gait, and why unsupervised / similarity-first", 2)
    r.p(
        "Gait kinematics are a high-dimensional, repeatable motor signature: every walking "
        "cycle yields joint-angle and marker trajectories that can, in principle, carry "
        "information about how a person organizes locomotion. The scientific temptation, "
        "given a victimization survey joined to a motion-capture archive, is to train a "
        "classifier and report accuracy. With n = 31 independent people and thousands of "
        "derived features, that approach can produce impressive in-sample accuracy that "
        "does not correspond to a shared, interpretable biomechanical difference. The "
        "project therefore deferred predictive modeling and asked first whether any "
        "difference is statistically supported at the subject level, robust to leave-one-person-out "
        "and multiplicity, and interpretable in anatomical or temporal terms."
    )
    r.p(
        "After Phases 3–6 returned null or non-enriching answers to the mean-difference / "
        "phenotype / compactness / trajectory questions, a second layer was still required. "
        "Victims might not be close to each other in Euclidean space, and no feature mean "
        "might survive FDR, while still sharing a relative pattern: a common direction of "
        "deviation, a shared set of “abnormal” features, a shared waveform shape after "
        "stripping amplitude, a phase-localized signature, or a shared inter-joint coupling. "
        "Similarity P0 was built to test those constructs on frozen Phases 0–6 artifacts, "
        "without rewriting those pipelines."
    )

    r.h("The independent-unit issue", 2)
    r.p(
        "The archive contains 880 usable gait cycles from 260 walking trials. Treating those "
        "cycles as 880 independent observations would be pseudo-replication: cycles from the "
        "same person are repeated measures, not new people. Inflating n from 31 to 880 would "
        "shrink p-values and produce false confidence. Every inferential step in Phases 3–6 "
        "and Similarity P0 therefore permutes or compares subjects. Cycle-level data are used "
        "to form a subject summary (typically a median or, for CRP, a circular mean) and then "
        "discarded as the inferential unit."
    )
    r.table(
        ["Quantity", "Role"],
        [
            ["31 subjects", "Independent sampling units for inference"],
            ["17 victimized / 14 control", "Group labels (Y / N)"],
            ["260 walking trials (WU*)", "Repeated sessions within people"],
            ["880 gait cycles (440 L / 440 R)", "Repeated measures — not 880 independent samples"],
        ],
        widths=[2.4, 4.1],
        caption="Table 2.1. Independent unit versus repeated measures (README §1; Phase 0/1 counts).",
    )

    r.h("Cohort and data capture", 2)
    r.p(
        "The processed MATLAB file contains females only. Phase 0 verified 31 subjects "
        "(expected 31), 17 victimized and 14 non-victimized, join key Excel “Subject No” ↔ "
        "MATLAB S#, 260 walking trials (expected 260), 242 valid walking trials, and 100 Hz "
        "sampling. The raw MAT has 43 subjects; 12 males are present in raw but not processed "
        "(S6, S10, S18, S21, S22, S25, S29, S33, S36, S45, S47, S50). Irregular session names "
        "(WU0, WU01Copy, and similar) were identified and not renamed."
    )
    r.p(
        "Subject IDs in the processed file are S2–S5, S7–S9, S11–S15, S17, S19, S23, S26, S27, "
        "S30–S32, S34, S35, S37–S43, S46, S48. Walking trials per subject range from 4 to 12 "
        "(median 9.0, sum 260): victims min 4 / median 9.0 / max 11 (sum 143); controls min 5 / "
        "median 9.0 / max 12 (sum 117). Phase 1 extracted 880 usable lower-body cycles "
        "(all PASS; 0 PASS WITH WARNINGS; 0 FAIL), 440 left and 440 right, time-normalized to "
        "101 points (0–100%). Cycle duration_seconds on the 880-row inventory has minimum 0.88 s, "
        "median 1.04 s, maximum 1.32 s (mean 1.044932 s, SD 0.069127 s), matching the README’s "
        "rounded “roughly 0.88–1.32 s (median ~1.04 s).” Usable cycles per subject are uneven: "
        "victims min 4, median 32.0, max 44, sum 487; controls min 3, median 30.5, max 44, sum 393. "
        "S19 has 4 usable cycles; S30 has 3."
    )
    r.p(
        "Anthropometry is not tabulated in the Phase 0/1 reports. The processed-dataset blueprint "
        "lists mass 47.7–100.0 kg and height 141.0–167.6 cm across the 31 females (S15 is the "
        "short outlier at 141.0 cm). README §2 rounds this to “mass ~48–100 kg; height ~141–168 cm”; "
        "this report uses the blueprint extrema. Capture is Plug-in Gait kinematics at 100 Hz. "
        "Spatial and angle axes are stored as ax1 / ax2 / ax3 and are not certified as "
        "anterior–posterior / mediolateral / vertical."
    )
    r.p(
        "Victimization is a survey label, not a controlled exposure. Among the 17 labeled Y, "
        "VictimType is Nd 7, in-person (Ip) 6, Both 3, online 1 (blueprint §3). CyberBullied "
        "overall: Yes 6, Nd 6, No 19 (No includes all 14 non-victims). Primary ML label remains "
        "binary Y/N; subtype-stratified inference was not a discovery test in Phases 3–6 or P0."
    )
    r.table(
        ["Item", "Value", "Source"],
        [
            ["Sex in processed file", "Female only", "Phase 0 / blueprint"],
            ["Subjects", "31 (PASS)", "Phase 0 audit_report.md"],
            ["Victimized (Y) / control (N)", "17 / 14 (PASS)", "Phase 0"],
            ["Walking trials WU*", "260 (PASS); 242 valid", "Phase 0"],
            ["Sampling rate", "100 Hz (PASS)", "Phase 0"],
            ["Usable gait cycles", "880 (440 L / 440 R), all PASS", "Phase 1"],
            ["Time base", "0–100%, 101 points", "Phase 1"],
            ["Cycle duration (s)", "min 0.88, median 1.04, max 1.32", "Phase 1 inventory"],
            ["Mass (kg)", "47.7–100.0", "blueprint (not Phase 0/1)"],
            ["Height (cm)", "141.0–167.6 (S15 = 141.0)", "blueprint"],
            ["VictimType among Y=17", "Nd 7, Ip 6, Both 3, online 1", "blueprint §3"],
            ["KinFC side encoding", "code 1 = right, code 2 = left", "Phase 1 (measured)"],
        ],
        widths=[2.2, 2.5, 1.8],
        caption="Table 2.2. Cohort and capture specifications.",
    )

    # ============================================================
    # 3. Pipeline
    # ============================================================
    r.h("3. Pipeline Overview", 1)
    r.p(
        "Work proceeds in two layers. Phases 0–6 ask whether victimized and non-victimized "
        "females differ on subject-level summaries, unsupervised phenotypes, Euclidean "
        "compactness, or time-resolved trajectories. After that layer was frozen, Similarity P0 "
        "asks whether victims share a pattern with each other that controls do not. P1 and "
        "Phase 7 were not started."
    )
    r.table(
        ["Stage", "What it does structurally"],
        [
            ["Raw MAT + survey", "43-subject archive and Excel labels; processed file is 31 females with survey joined on Subject No."],
            ["Phase 0", "Read-only audit: counts, join, events, missingness, irregularities. Does not extract cycles or features."],
            ["Phase 1", "Parse KinFC/KinFO/Midsvnt; extract ipsilateral FC-to-FC cycles; normalize core signals to 101 points."],
            ["Phase 2", "Label-blind feature discovery at cycle level and median/mean/std/cv aggregation to subjects."],
            ["Phase 3", "Label-blind screen and redundancy, then 17 vs 14 tests with FDR, effect size, LOSO, consistency."],
            ["Phase 4", "Label-blind family-PC clustering; labels joined only after assignments freeze; then enrichment."],
            ["Phase 5", "Within-victim Euclidean similarity and subgroup search in the frozen Phase 4 27-D space."],
            ["Phase 6", "Subject-median trajectories; cluster permutation of subject labels; shape and asymmetry secondaries."],
            ["★ Freeze", "Phases 0–6 pipelines and certified outputs are not rewritten by later work."],
            ["P0.1", "Shared deviation direction (mean pairwise cosine of control-referenced 27-D vectors)."],
            ["P0.2", "Shared abnormality set (Jaccard of binary exceedances on 30 locked features)."],
            ["P0.3", "Shared waveform shape after z-score (Pearson and DTW on 12 locked curves)."],
            ["P0.4", "Event-window localization (5 reconstructable phases × 12 curves × 2 aggregations × 2 tests = 240 FDR)."],
            ["P0.5", "Confound residualization folded into each P0 test (not a separate discovery script)."],
            ["P0.6", "Shared Hilbert CRP coupling on 6 locked pairs (circular similarity and DTW)."],
            ["Power / MDE", "Post-hoc P0.1 simulation: how large a shared direction n=17 vs 14 could detect."],
            ["P1 / Phase 7", "Not started. No Wasserstein/RV/soft-DTW/common subspace; no predictive ML."],
        ],
        widths=[1.5, 5.0],
        caption="Table 3.1. Structural flow (README §3), one sentence per stage.",
    )
    r.p(
        "Hard rules across both layers: victimization labels are not used to build features, "
        "scale, PCA, choose clusters, or pick which curves/pairs/phases to test (those lists "
        "are pre-registered before looking at group results). Labels are used for victim/control "
        "tests and for similarity among the labeled victim subset only after the relevant "
        "representation is frozen. The program does not manufacture a signature if FDR, "
        "robustness, or consistency fail, and does not train classifiers “to prove” a gait exists."
    )

    # ============================================================
    # 4. Walkthrough
    # ============================================================
    r.h("4. Phase-by-Phase Detailed Walkthrough", 1)
    r.p(
        "Each subsection below uses the same seven headings: Question, Input, Method, "
        "Pre-registration, Result, Decision, Output artifacts. Numbers are from the cited "
        "primary files. Inferential seed, unless a subsection says otherwise, is 20260813."
    )

    # ---- Phase 0 ----
    r.step(
        "4.1 Phase 0 — Dataset audit",
        "What, exactly, is in the processed motion-capture archive, how does it join to the "
        "victimization survey, and is the file complete enough to support subject-level gait "
        "analysis — without modifying any source data?",
        "Read-only access to data/processed/Data_structure_all_subs.mat (31 females), "
        "data/raw/Data_structure_all_subs.mat (43 subjects), and data/raw/Victimization surveys.xlsx. "
        "Phase 0 consumes no prior analysis artifact; it is the first computational step.",
        [
            "The audit is a structured inventory, not a statistical test. It counts subjects, "
            "walking trials (WU*), sampling rate, markers (expected 37), joint angles (expected 26), "
            "joint centers (6), COM fields, and gait events KinFC / KinFO / Midsvnt on every walking "
            "trial. It verifies the join Excel Subject No ↔ MATLAB S# and the 17/14 label split. "
            "It flags NaNs, missing markers, irregular session names, and trial imbalance. "
            "Critical issues would block later phases; warnings document quality without renaming "
            "or imputing. Source files are not written.",
        ],
        "Phase 0 has no scientific pre-registration of a test family. Expected counts (31 subjects, "
        "260 WU* trials, 17/14 labels, 100 Hz, 37 markers, 26 joint angles) are the audit’s "
        "pass/fail criteria, stated in the report’s check list.",
        [
            "Status is PASS WITH WARNINGS, generated 2026-08-13. Subjects 31/31 PASS. Walking "
            "trials 260/260 PASS; valid walking trials 242. Sampling 100 Hz PASS. Labels 17 Y / 14 N, "
            "split check PASS, join PASS. Events: KinFC 260/260, KinFO 260/260, Midsvnt 260/260. "
            "Markers 37, joint angles 26, modal observed kinematics count 86 PASS. Raw MAT has 43 "
            "subjects; 12 IDs present in raw but not processed (males). Critical issues: 0. "
            "Warnings: 74, dominated by upper-arm and static-trial NaNs (LUPA/RUPA), S12 Copy session "
            "names, S17 non-canonical WU names, S30 missing LUPA on four trials (85 kinematics fields "
            "versus 86), and trial imbalance (4–12 walking trials per subject, median 9.0). "
            "Arm-marker gaps do not, by later Phase 1 policy, invalidate lower-body cycles.",
        ],
        "Table 4.1. Phase 0 audit checks (results/phase0/audit_report.md).",
        ["Check", "Result"],
        [
            ["subject_count", "PASS (31)"],
            ["label_split", "PASS (17 Y / 14 N)"],
            ["join (Subject No ↔ S#)", "PASS"],
            ["sampling_rate", "PASS (100 Hz)"],
            ["walking_trials", "PASS (260 WU*; 242 valid)"],
            ["walking_signal_count", "PASS (modal 86)"],
            ["events_KinFC / KinFO / Midsvnt", "PASS (260/260 each)"],
            ["raw_mat / survey Excel / survey table", "PASS"],
            ["Critical issues", "0"],
            ["Warnings", "74"],
        ],
        [3.4, 3.1],
        [
            "Decision: PASS WITH WARNINGS. The processed file is verified, the survey join holds, "
            "events exist on all walking trials, and no source data were modified. Warnings are "
            "quality annotations (especially upper-body missingness and irregular names), not a "
            "failed cohort. Gait-cycle min/median/max per subject is deferred to Phase 1.",
        ],
        [
            "results/phase0/audit_report.md",
            "docs/phase0_dataset_audit.md (companion write-up)",
            "Code: src/gait_research/ (catalog.py, matio.py, sessions.py, audit.py, labels.py); scripts/audit_dataset.py",
        ],
    )

    # ---- Phase 1 ----
    r.step(
        "4.2 Phase 1 — Gait events and cycles",
        "Can ipsilateral foot-contact-to-contact gait cycles be extracted from every walking "
        "trial, with empirically validated left/right encoding, and can core lower-body signals "
        "be time-normalized to a common 0–100% grid without modifying the processed MAT?",
        "The Phase 0–verified processed MAT and its KinFC, KinFO, and Midsvnt event tables on "
        "all 260 walking trials. The processed MAT remains the canonical trajectory store; "
        "normalized arrays are written under results/phase1/ only.",
        [
            "A gait cycle is ipsilateral foot contact to the next ipsilateral foot contact. "
            "KinFC column 2 is a side code. Heel-Z at each contact (lower heel = contacting foot) "
            "across all 260 walking trials maps code 1 → right and code 2 → left. This was "
            "measured, not assumed; a user-facing example that treated 1 as left is incorrect "
            "for this dataset. Mapping mismatches versus the heel vote: 0. Event validation "
            "PASS/WARNING/FAIL = 260/0/0.",
            "Quality policy: missing opposite foot-contact or implausible duration outside "
            "0.50–2.20 s fails a cycle. Missing FO, 0 or >2 mid-stance events, unusual duration "
            "outside 0.70–1.60 s, or incomplete lower-body yields PASS WITH WARNINGS but remains "
            "usable for lower-body gait. Upper-arm gaps (LUPA/RUPA/RFRM) do not fail a cycle if "
            "lower-body coverage is intact. Upper-body gaps are recorded per domain and do not "
            "change overall cycle status. Core signals are interpolated onto 101 points (0–100%). "
            "LPelvisAngles / RPelvisAngles are not in the Phase 1 core cube; pelvis is represented "
            "by markers (LASI, RASI, LPSI, RPSI). Provenance: cycle_id encodes Subject_Trial_Side_Index; "
            "start_frame / end_frame are MATLAB 1-based KinFC frames.",
        ],
        "Phase 1 does not pre-register a victim/control test. The core-signal list and cycle "
        "definition are the locked extraction contract for all later phases. Similarity P0.4 "
        "later audited that all 880 cycles store IC, opposite FO, mid-stance, opposite FC, "
        "ipsilateral FO, and next IC in strict order (100% complete).",
        [
            "Status PASS, generated 2026-08-13. Total gait cycles 880; left 440; right 440; "
            "PASS 880; PASS WITH WARNINGS 0; FAIL 0; usable for lower-body 880; normalized 880. "
            "Victims (n=17): min 4, median 32.0, max 44, sum 487 usable cycles. Controls (n=14): "
            "min 3, median 30.5, max 44, sum 393. Lowest counts: S19 (victim) 4 cycles from 4 trials; "
            "S30 (control) 3 cycles from 5 trials. Inventory duration_seconds: min 0.88, median 1.04, "
            "max 1.32. Normalized core cube: results/phase1/gait_cycles/normalized_core.npz with "
            "shape (880, 26, 101, 3) as documented in README; 26 core signals listed in the Phase 1 report.",
        ],
        "Table 4.2. Phase 1 cycle extraction (phase1_report.md and gait_cycle_inventory.csv).",
        ["Quantity", "Value"],
        [
            ["Walking trials inspected", "260 / 260"],
            ["Event validation PASS / WARNING / FAIL", "260 / 0 / 0"],
            ["Heel-vote mapping mismatches", "0"],
            ["Usable cycles (L / R)", "880 (440 / 440)"],
            ["Cycle quality FAIL or WARN", "0 / 0"],
            ["Victim cycle min / median / max / sum", "4 / 32.0 / 44 / 487"],
            ["Control cycle min / median / max / sum", "3 / 30.5 / 44 / 393"],
            ["duration_seconds min / median / max", "0.88 / 1.04 / 1.32"],
            ["KinFC side", "1 = R, 2 = L (measured)"],
        ],
        [3.4, 3.1],
        [
            "Decision: PASS. All walking trials yielded usable lower-body cycles; side encoding "
            "is empirically validated; source data were not modified. Uneven cycle counts (especially "
            "S19 and S30) remain a limitation for subject medians, not a failed extraction.",
        ],
        [
            "results/phase1/phase1_report.md",
            "results/phase1/gait_cycle_inventory.csv (880 rows)",
            "results/phase1/gait_cycles/normalized_core.npz",
            "scripts/extract_gait_cycles.py",
        ],
        extra_prose=[
            "Core gait signals normalized (Phase 1 list): LASI, RASI, LPSI, RPSI, LHJC, RHJC, "
            "LHipAngles, RHipAngles, LKJC, RKJC, LKneeAngles, RKneeAngles, LAJC, RAJC, LAnkleAngles, "
            "RAnkleAngles, LAbsAnkleAngle, RAbsAnkleAngle, LHEE, RHEE, LTOE, RTOE, LFootProgressAngles, "
            "RFootProgressAngles, CentreOfMass, CentreOfMassFloor.",
        ],
    )

    # ---- Phase 2 ----
    r.step(
        "4.3 Phase 2 — Feature discovery",
        "What label-blind, anatomically traced gait features can be computed on each of the "
        "880 cycles and aggregated to one row per subject, so that later statistics never "
        "invent features after seeing victimization labels?",
        "Phase 1 normalized core signals (101-point cycles) and the cycle inventory (durations, "
        "events, side). Victimization labels are not inputs to feature construction or aggregation.",
        [
            "Cycle-level families cover kinematics (ROM, mean, SD, peaks, timing, derivatives), "
            "temporal event percentages, spatial path lengths, 10% phase bins, coordination "
            "(correlation and lag), and smoothness. Derivatives use Savitzky–Golay window 11, "
            "polynomial 3 on the 101-point cycle; ROM/min/max are unsmoothed. Spatial axes remain "
            "ax1/ax2/ax3. Subject aggregation default is the median across that subject’s cycles; "
            "the subject table also stores mean, std, cv, and n for each cycle feature, plus "
            "ipsilateral symmetry (L-cycle left limb versus R-cycle right limb, not same-window "
            "pooling of all left versus all right) and variability features. No silent imputation "
            "of ROM (finite-only drop); derivative-path NaN interpolation is catalog-documented. "
            "No victim-versus-control test is run here.",
        ],
        "There is no group-test lock file. The feature catalog (805 specs) and the rule that "
        "labels_used = False are the Phase 2 contract. Certification verifies no label leakage "
        "in tables or source tokens, catalog-versus-table reconciliation (714 cycle specs + 91 "
        "subject extras = 805), units vocabulary, anatomical metadata, phase bins 0–10 … 90–100, "
        "and that dropping another subject’s cycles does not change S3’s median — no cross-subject influence.",
        [
            "Generated 2026-08-13. Cycles 880; subjects 31; cycle-level feature columns 714; "
            "subject-level columns 3665; catalog entries 805; label columns leaked: none. "
            "Features with complete subject coverage 714; with some missingness 0; unavailable 0. "
            "Certification status PASS. Catalog families: kinematic 390, phase 240, spatial 64, "
            "symmetry 58, variability 33, temporal 9, coordination 8, smoothness 3. "
            "Recomputed check: S14 LKneeAngles_ax1_rom median = 65.71893581748009 matches stored. "
            "S3 same feature median 68.31004357337952 is unchanged after dropping S2 cycles.",
        ],
        "Table 4.3. Phase 2 coverage (phase2_report.md and phase2_certification.md).",
        ["Quantity", "Value"],
        [
            ["Cycle rows × feature columns", "880 × 714"],
            ["Subject rows × columns", "31 × 3665"],
            ["Catalog entries (cycle + extra)", "805 (714 + 91)"],
            ["Label leakage", "none (tables and code tokens)"],
            ["Savitzky–Golay", "window 11, poly 3"],
            ["Default subject summary", "*__median"],
            ["Certification", "PASS"],
        ],
        [3.4, 3.1],
        [
            "Decision: PASS (certified). Features exist, are anatomically traced, do not leak "
            "labels, and are ready for Phase 3 to default to *__median unless a dispersion feature "
            "is the scientific target. No group difference is claimed or tested.",
        ],
        [
            "results/phase2/phase2_report.md",
            "results/phase2/phase2_certification.md",
            "results/phase2/cycle_features.parquet, subject_features.parquet, feature_catalog.json",
            "scripts/run_phase2.py, scripts/certify_phase2.py",
        ],
    )

    r.doc.add_page_break()

    # ---- Phase 3 ----
    r.step(
        "4.4 Phase 3 — Statistical gait signature discovery",
        "After label-blind screening and redundancy reduction, which subject-level features, "
        "if any, show a victim-versus-control difference that survives FDR, is at least medium "
        "by Cliff’s δ, is directionally stable under leave-one-subject-out, and is consistent "
        "across victims — the pre-specified definition of a gait signature?",
        "Phase 2 subject_features.parquet and feature_catalog.json. Analysis columns: 805 "
        "(*__median, var_*, sym_*). Screening and redundancy used no group labels. Labels were "
        "joined only for the group comparison.",
        [
            "Quality screen (no victimized column allowed in the screened table) left 743 columns. "
            "Spearman |ρ| ≥ 0.90 clustering reduced these to 335 redundancy representatives. "
            "On those 335, a two-sided Mann–Whitney test produced raw p-values; Benjamini–Hochberg "
            "FDR was applied to that family of 335. Effect size is Cliff’s δ with subject-resampled "
            "bootstrap CIs. Permutation: 999 subject-label shuffles, seed 20260813; the permutation "
            "unit is the subject, never the cycle. LOSO direction agreement is the fraction of "
            "single-subject deletions that preserve the sign of the group difference. Victim "
            "directional consistency is the share of victims on the group-difference side of the "
            "control median. No classifier, victim score, or accuracy was computed.",
            "Pre-specified signature rule (all required): BH FDR q ≤ 0.10; |Cliff’s δ| ≥ 0.33 "
            "(medium); LOSO direction agreement ≥ 0.80; victim directional consistency ≥ 0.60. "
            "Ranking of an exploratory list uses effect magnitude, FDR weight, LOSO, consistency, "
            "coverage, and family interpretability — not p-value alone.",
        ],
        "The signature rule and the 335-test FDR family are specified in phase3_report.md before "
        "the ranked list. There is no separate JSON lock of which 335 columns would be tested "
        "beyond the label-blind screen/redundancy procedure. Labels are forbidden in screening "
        "(certification: screening_rejects_labels PASS).",
        [
            "Generated 2026-08-13. FDR ≤ 0.05: 0. FDR ≤ 0.10: 0. Signature-rule features: 0. "
            "Smallest Mann–Whitney raw p: 0.01104 in the report (multiple_testing.csv: 0.011039235074033062 "
            "on LKneeAngles_ax3_tmax_pct__median). Smallest subject-permutation p: 0.008 (same feature). "
            "Minimum FDR q in multiple_testing.csv is 0.9960585779461092; 0 features have q ≤ 0.10 or "
            "q ≤ 0.05. Uncorrected permutation p-values can look small and must not be read as a "
            "signature. The ranked list is exploratory. Knee region has the largest max |δ| among "
            "anatomical summaries (0.542) with 0 FDR pass.",
        ],
        "Table 4.4a. Phase 3 inferential gates (phase3_report.md; multiple_testing.csv n=335).",
        ["Gate", "Result"],
        [
            ["Analysis columns", "805"],
            ["Passed quality screen", "743"],
            ["Redundancy representatives (FDR family)", "335"],
            ["Permutations / seed / unit", "999 / 20260813 / subject"],
            ["FDR q ≤ 0.05", "0"],
            ["FDR q ≤ 0.10", "0"],
            ["Signature-rule hits", "0"],
            ["Smallest MW raw p", "0.011039 (report 0.01104)"],
            ["Smallest permutation p", "0.008"],
            ["Minimum FDR q", "0.9960585779461092"],
        ],
        [3.4, 3.1],
        [
            "Decision: COMPLETE / certification PASS WITH WARNINGS. The warning "
            "no_manufactured_signature is explicit: signature_rule hits = 0; exploratory ranks "
            "are not claimed as a confirmed signature. Interpretation in the source report: no "
            "feature met the signature rule; n=17 vs 14 has low power and a true medium effect can "
            "fail FDR; uncorrected small p-values are not multiplicity-controlled. Phase 4 "
            "independent validation was not run from this list as a confirmed set.",
        ],
        [
            "results/phase3/phase3_report.md",
            "results/phase3/phase3_certification.md",
            "results/phase3/candidate_signature.csv",
            "results/phase3/statistics/multiple_testing.csv",
            "scripts/run_phase3.py",
        ],
        extra_tables=[
            (
                "Table 4.4b. Exploratory top 10 (candidate_signature.csv; all status = exploratory, FDR q = 0.9961).",
                ["Rank", "Feature", "Dir.", "Cliff δ", "perm p", "LOSO", "Vic. cons."],
                [
                    ["1", "LHipAngles_ax3_tpeak_vel_pct__median", "LOWER", "−0.4916", "0.028", "1.00", "0.882"],
                    ["2", "LKneeAngles_ax3_tmax_pct__median", "LOWER", "−0.5420", "0.008", "1.00", "0.765"],
                    ["3", "RKneeAngles_ax3_peak_acc__median", "HIGHER", "0.4874", "0.024", "1.00", "0.824"],
                    ["4", "LFootProgressAngles_ax2_mean__median", "LOWER", "−0.4370", "0.047", "1.00", "0.882"],
                    ["5", "LFootProgressAngles_ax1_std__median", "LOWER", "−0.4118", "0.058", "1.00", "0.882"],
                    ["6", "RAnkleAngles_ax3_peak_acc__median", "HIGHER", "0.39496", "0.062", "1.00", "0.882"],
                    ["7", "LHipAngles_ax2_tpeak_vel_pct__median", "LOWER", "−0.4244", "0.054", "1.00", "0.765"],
                    ["8", "LAnkleAngles_ax1_phase_40_50_rom__median", "LOWER", "−0.4622", "0.037", "1.00", "0.765"],
                    ["9", "RFootProgressAngles_ax3_peak_acc__median", "HIGHER", "0.4118", "0.070", "1.00", "0.765"],
                    ["10", "LKneeAngles_ax3_peak_acc__median", "HIGHER", "0.3529", "0.101", "1.00", "0.882"],
                ],
                [0.55, 2.55, 0.75, 0.7, 0.55, 0.55, 0.85],
            )
        ],
    )

    # ---- Phase 4 ----
    r.step(
        "4.5 Phase 4 — Phenotypes and heterogeneity",
        "Do the 31 subjects form stable, interpretable gait phenotypes in a label-blind "
        "representation, and only then is victimization disproportionately represented in any "
        "phenotype? This is a heterogeneity question, not a second attempt to force a "
        "population-wide mean signature after Phase 3’s null.",
        "Phase 2 subject features restricted to the 335 Phase 3 label-blind redundancy "
        "representatives. Independent unit n = 31 subjects; 880 cycles are not clustering units. "
        "Phase 1 normalized cycles are used only later for phenotype-median trajectory plots, "
        "after dropping inventory victimization columns.",
        [
            "Victimization labels were absent from feature selection, median/IQR scaling, family "
            "PCA, global PCA, clustering, k selection, stability, and phenotype characterization. "
            "Labels were joined only after assignments were frozen. Compact representation: "
            "within-family PCA keeping ≥80% family variance, cap 8 PCs per family, then divide "
            "by √n_pcs so families do not dominate Euclidean distance, yielding 27 dimensions "
            "(coordination 6→2, kinematic 163→4, phase 71→8, smoothness 2→1, spatial 21→5, "
            "symmetry 44→3, temporal 3→2, variability 25→2). Scaling: median/IQR on all 31 "
            "subjects without labels; non-finite values imputed with the column median; zero-IQR "
            "columns set to 0. PC1 explains 37.2% of compact-space variance; components were not "
            "chosen to separate victims and controls. Clustering used the family-PC matrix.",
            "Primary algorithm: hierarchical Ward. Sensitivity: k-means, 10 random inits, seed "
            "20260813. Candidate k ∈ {2,3,4}. Selection used silhouette, minimum cluster size ≥ 4, "
            "and subject-bootstrap ARI (150 replicates of 80% subject subsamples); victim/control "
            "separation was not a criterion. Stability also includes leave-one-subject-out ARI. "
            "Enrichment: Fisher exact and subject-label permutation, FDR across phenotypes. "
            "Confounding: Kruskal–Wallis of mass, height, left and right leg length versus phenotype.",
        ],
        "The stability rule is pre-specified in the selection payload: min_cluster_size 4, "
        "min_silhouette 0.2, min_mean_boot_ari 0.5. k is not chosen from labels "
        "(certification k_not_from_labels PASS). Assignment table has no victimized column.",
        [
            "Generated 2026-08-13. Selected k = 2, method hierarchical, reason "
            "max_bootstrap_ari_among_stable_hierarchical. Silhouette = 0.41503306618271296 "
            "(README rounds to 0.42). Mean bootstrap ARI = 0.7546071898190309. Minimum size = 4. "
            "Sizes: phenotype 1 n = 27, phenotype 2 n = 4. Phenotype 2 members (phenotype_assignments.csv): "
            "S5, S19, S35, S40. After labels: phenotype 1 is 15/27 victimized (prop 0.56, expected 0.55, "
            "Fisher p = 1, perm p = 1, FDR q = 1); phenotype 2 is 2/4 victimized (prop 0.50, expected 0.55, "
            "same p = 1). Height Kruskal p = 0.0499; mass p = 0.976; lleg p = 0.882; rleg p = 0.836.",
        ],
        "Table 4.5. Phase 4 clustering and enrichment (phase4_report.md, certification, assignments).",
        ["Quantity", "Value"],
        [
            ["Compact dimensions", "27 family PCs from 335 representatives"],
            ["k / method", "2 / hierarchical Ward"],
            ["Silhouette", "0.41503306618271296"],
            ["Mean bootstrap ARI (150 × 80% subjects)", "0.7546071898190309"],
            ["Cluster sizes", "27 vs 4"],
            ["Outgroup IDs", "S5, S19, S35, S40"],
            ["Phenotype 1 victimized", "15/27 (expected 0.55); Fisher p=1; perm p=1"],
            ["Phenotype 2 victimized", "2/4 (expected 0.55); Fisher p=1; perm p=1"],
            ["Height Kruskal p", "0.0499"],
            ["Mass / L leg / R leg Kruskal p", "0.976 / 0.882 / 0.836"],
        ],
        [3.2, 3.3],
        [
            "Decision: PASS WITH WARNINGS. Certification warnings: not_two_large_phenotypes "
            "(majority/outgroup split — do not over-interpret as two gait types) and "
            "height_association (phenotype may partly track stature). Source interpretation: "
            "the n=4 group should not be promoted to a named clinical phenotype. Composition is "
            "compatible with the overall 17/31 victim base rate. This is not a claim that "
            "victimization causes gait and is not a predictive model.",
        ],
        [
            "results/phase4/phase4_report.md",
            "results/phase4/phase4_certification.md",
            "results/phase4/phenotype_assignments.csv (no victim column)",
            "scripts/run_phase4.py",
        ],
    )

    # ---- Phase 5 ----
    r.step(
        "4.6 Phase 5 — Within-victim Euclidean structure",
        "Are the 17 victimized subjects more similar to each other in the certified gait "
        "representation than chance, and do they form stable gait subgroups that differ from "
        "the 14 controls — a question Phase 3 (means) and Phase 4 (whole-cohort phenotypes) "
        "did not close?",
        "The certified Phase 4 family-balanced compact matrix (31 × 27), built without "
        "victimization in scaling, PCA, or feature selection. Labels are used only to subset "
        "victims and to test similarity/enrichment. Independent unit: subject (n=31; 17 victimized). "
        "Cycles are not clustering units.",
        [
            "Within-victim similarity: mean pairwise Euclidean distance among the 17 victims in "
            "the 27-D space. Null: 999 permutations drawing random groups of 17 from 31, unit = "
            "subject, alternative = group more similar than a random same-size group (so smaller "
            "distance is the “similar” tail). Nearest neighbors: fraction of victims whose 1-NN "
            "is another victim, with a matching permutation; also mean 3-NN victim fraction. "
            "Victim subgroups: hierarchical clustering among victims with the same stability "
            "thresholds as Phase 4 (min size 4, min silhouette 0.2, min mean bootstrap ARI 0.5), "
            "k chosen from those metrics, not from victim-versus-control separation. All-17-versus-14 "
            "centroid distance is permutation-tested. If no stable subgroup exists, subgroup-versus-control "
            "and subgroup-versus-other-victim characterizations are not interpreted.",
        ],
        "Stability criteria are the same payload as Phase 4. Certification requires permutation "
        "unit = subject, assignment table without a victimized column, and an honest_no_subgroup "
        "warning if no split is forced.",
        [
            "Generated 2026-08-13. Observed mean pairwise distance among 17 victims: 20.41683891912052. "
            "Permutation null mean 20.06005494755659 (SD 2.4785971272835554; 5th–95th "
            "15.845444901593474–24.145583593455807). Permutation p = 0.547 (n_perm = 999). "
            "Victim–victim / control–control / victim–control mean pairwise: 20.41683891912052 / "
            "20.0157339007024 / 19.75387534385816. Victims are slightly farther from each other "
            "than the null mean, and victim–control pairs are on average closer than victim–victim "
            "pairs — the opposite of a tight victim neighborhood. 1-NN victim fraction 0.5294117647058824 "
            "versus null mean 0.5361832420655951, perm p = 0.615; mean 3-NN victim fraction "
            "0.4705882352941176. Selected k = None; reason no_stable_phenotype_structure; number of "
            "stable subgroups 0. All 17 vs 14 centroid distance 4.5893, perm p = 0.771.",
        ],
        "Table 4.6. Phase 5 similarity (phase5_report.md; within_victim_similarity.csv; nn_permutation.csv).",
        ["Statistic", "Value"],
        [
            ["Victim mean pairwise distance", "20.41683891912052"],
            ["Null mean (random 17 of 31)", "20.06005494755659"],
            ["Null SD / p05 / p95", "2.4786 / 15.8454 / 24.1456"],
            ["Similarity permutation p (n=999)", "0.547"],
            ["Control–control mean pairwise", "20.0157339007024"],
            ["Victim–control mean pairwise", "19.75387534385816"],
            ["1-NN is another victim", "0.5294 (null 0.5362, p=0.615)"],
            ["Mean 3-NN victim fraction", "0.4706"],
            ["Centroid distance 17 vs 14 / p", "4.5893 / 0.771"],
            ["Stable victim subgroups", "0 (k=None)"],
        ],
        [3.4, 3.1],
        [
            "Decision: PASS WITH WARNINGS (honest_no_subgroup). Source conclusion: no robust "
            "within-victim gait structure was detected. Victims are not more similar to each other "
            "than random groups of 17 in the certified representation, or any candidate split failed "
            "stability/size rules. This does not revive a population-wide victim signature. "
            "Phase 6 had not been started when this report was written; it was run subsequently "
            "as the next frozen phase.",
        ],
        [
            "results/phase5/phase5_report.md",
            "results/phase5/phase5_certification.md",
            "results/phase5/similarity/within_victim_similarity.csv",
            "results/phase5/neighbors/nn_permutation.csv",
            "scripts/run_phase5.py",
        ],
    )

    # ---- Phase 6 ----
    r.step(
        "4.7 Phase 6 — Time-resolved trajectories",
        "Do victimized and non-victimized females differ in normalized gait trajectories "
        "(0–100%, 101 points) after subject-level aggregation — i.e., did collapsing cycles "
        "into Phase 2/3 scalars hide localized time-resolved differences?",
        "Phase 1 certified normalized_core.npz (880 cycles, 101 points). No renormalization. "
        "Inferential n is 31 subjects, never 880. 86 channels were tested; quality control "
        "required all 31 subjects to have ≥90% finite time points (ineligible channels n = 0). "
        "No silent zero-fill; no inferential interpolation.",
        [
            "Within-subject nanmedian across that subject’s cycles (mean stored as sensitivity); "
            "no pooling across people. Primary statistic: Welch t-statistic time series; "
            "cluster-forming threshold |t| > 2.045; cluster mass = sum of |t| on contiguous "
            "suprathreshold points. Cluster-based permutation of subject labels (each 101-point "
            "trajectory kept intact; never shuffle time or cycles) uses max cluster mass to "
            "control the family of 101 time points within a channel. Primary permutations = 9999; "
            "secondary = 1999; seed = 20260813. Benjamini–Hochberg FDR is applied within analysis "
            "level (primary / primary_asymmetry / secondary / secondary_asymmetry).",
            "Primary channels (frozen): LHipAngles_ax1, RHipAngles_ax1, LKneeAngles_ax1, "
            "RKneeAngles_ax1, LAnkleAngles_ax1, RAnkleAngles_ax1, LFootProgressAngles_ax1, "
            "RFootProgressAngles_ax1, CentreOfMass_ax1, CentreOfMass_ax2, CentreOfMass_ax3. "
            "Shape analysis: Savitzky–Golay 11/3 as in Phase 2; peak/min timing and magnitude, "
            "n extrema, velocity RMS; Mann–Whitney + BH — none FDR ≤ 0.10. Bilateral asymmetry: "
            "A(t) = L−R and |L−R| on hip, knee, ankle, foot-progression ax1. Consistency: share of "
            "victims on the group-difference side of the control median at each time point. LOSO: "
            "sign of the regional mean difference after dropping each subject. Bootstrap: 1000 "
            "within-group subject resamples; percentile CI for regional mean median-difference.",
            "Region classes: ROBUST requires the predefined primary channel, cluster permutation "
            "and FDR, medium-or-larger effect, consistency, LOSO sign agreement, bootstrap CI "
            "excluding 0, and a contiguous phase span. EXPLORATORY is a weaker pre-specified "
            "class (cluster p < 0.05 in the report’s language). UNSUPPORTED is reported for "
            "completeness and is not a finding.",
        ],
        "Primary channel list is frozen (certification primary_predefined PASS). Axes remain "
        "ax1/ax2/ax3. Cluster threshold |t| = 2.045 is in phase6_config.json.",
        [
            "Generated 2026-08-13. ROBUST = 0; EXPLORATORY = 0. Strongest (lowest permutation p) "
            "region: RHipAngles_ax1 77.0–83.0%, δ = 0.42617046818727494, permutation p = 0.1469, "
            "q = 0.1469, class = UNSUPPORTED, bootstrap CI (−2.756730256761823, 17.708649499927247) "
            "includes 0, LOSO stability 1.0, analysis_level = primary, n_perm = 9999, unit = subject. "
            "Shape analysis: none FDR ≤ 0.10. Outcome A in the source report: no robust time-resolved "
            "victim-associated gait difference was detected.",
        ],
        "Table 4.7. Phase 6 trajectory inference (phase6_report.md; candidate_trajectory_regions.csv).",
        ["Quantity", "Value"],
        [
            ["Channels tested / excluded for missingness", "86 / 0"],
            ["Time points", "101"],
            ["Aggregation", "subject nanmedian"],
            ["Cluster |t| threshold", "2.045"],
            ["Primary / secondary permutations", "9999 / 1999"],
            ["ROBUST regions", "0"],
            ["EXPLORATORY regions", "0"],
            ["Strongest unsupported", "RHipAngles_ax1 77.0–83.0%"],
            ["That region δ / perm p / q", "0.42617 / 0.1469 / 0.1469"],
            ["That region bootstrap CI", "[−2.757, 17.709] includes 0"],
            ["Shape FDR ≤ 0.10", "none"],
        ],
        [3.4, 3.1],
        [
            "Decision: PASS WITH WARNINGS (honest_null_allowed: robust_findings = 0). "
            "Aggregation into Phase 2/3 scalars was not hiding a clear localized trajectory effect. "
            "Recommendation in the source report: do not train a victim classifier on these 31 people; "
            "if any exploratory region is pursued, preregister it on new subjects. Phase 7 was not started.",
        ],
        [
            "results/phase6/phase6_report.md",
            "results/phase6/phase6_certification.md",
            "results/phase6/phase6_config.json",
            "results/phase6/candidate_trajectory_regions.csv",
            "scripts/run_phase6.py",
        ],
    )

    r.doc.add_page_break()
    fill_p0_and_rest(r)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(OUT))
    print("Wrote", OUT)


def fill_p0_and_rest(r: Report):
    """P0 battery, power, cross-cutting, synthesis, limitations, appendix."""

    r.h("4.8 P0.1 — Deviation-direction alignment", 2)
    r.h3_block("a", "Question")
    r.p(
        "Do the 17 victims share a common direction of deviation from the control centroid in "
        "Phase 4’s 27-D family-PC gait space — a relative-pattern question that Euclidean "
        "compactness (Phase 5) can miss if victims lie along a common ray but not a tight ball?"
    )
    r.h3_block("b", "Input")
    r.p(
        "Frozen Phase 4 compact representation (31 × 27 family PCs) loaded via the similarity "
        "package. Subject-level covariates for residualization: height_cm, mass_kg, mean_leg_cm, "
        "cycle_duration_s_median. Phases 0–6 files are read, not rewritten. Generated 2026-08-16."
    )
    r.h3_block("c", "Method")
    r.p(
        "For each subject, d_i = x_i − mean(controls). The primary statistic is the mean pairwise "
        "cosine among the 17 victim d_i (one-sided greater: more alignment than chance). Null: "
        "9999 subject-label shuffles (choose 17 of 31 without replacement each permutation), "
        "seed 20260813. Bootstrap 95% CI on the observed cosine; mean cosine of victims toward "
        "the victim-mean direction; fraction of pairwise cosines > 0; LOSO sign agreement on the "
        "statistic. Residualization: OLS of each of the 27 columns on the four covariates with "
        "intercept, across subjects, before recomputing d_i. Primary evidence is the post-residual "
        "result. Gate for a defensible shared-direction signal: permutation p ≤ 0.05, LOSO pass, "
        "and bootstrap CI excluding ≤0 after residualization."
    )
    r.h3_block("d", "Pre-registration")
    r.p(
        "The construct (control-referenced cosine in the frozen 27-D space) and the gate are "
        "stated in p01_report.md. There is no separate feature-list JSON; the representation is "
        "the already-certified Phase 4 matrix. Residual covariates are the project-wide four."
    )
    r.h3_block("e", "Result")
    r.p(
        "Pre-residual mean pairwise cosine among victims is 0.06399477101002826 (report table "
        "0.0640) versus null mean 0.18945611077177135, permutation p = 0.9164 — observed below "
        "the null. Post-residual (primary) cosine is 0.05181913119541501 (report 0.0518; README "
        "and p0_synthesis round to 0.052) versus null mean 0.11029207303487745, permutation p = 0.7579 "
        "(synthesis 0.758). Bootstrap CI [0.0018289451414393416, 0.2512789721003968] includes 0. "
        "LOSO sign agreement 1.000, pass True. Consistency (fraction of pairwise cosines > 0) 0.882. "
        "Victims are less directionally aligned than random groups of 17."
    )
    r.table(
        ["Metric", "Pre-residual", "Post-residual (primary)"],
        [
            ["Mean pairwise cosine", "0.06399477101002826", "0.05181913119541501"],
            ["95% bootstrap CI", "[0.0163, 0.2754]", "[0.0018, 0.2513]"],
            ["Permutation p (greater)", "0.9164", "0.7579"],
            ["Null mean / 95th pct", "0.1895 / 0.3652", "0.1103 / 0.2566"],
            ["Mean cosine → victim-mean dir.", "0.2339", "0.3099"],
            ["Consistency (frac cosines > 0)", "0.941", "0.882"],
            ["LOSO sign agree / pass", "1.000 / True", "1.000 / True"],
            ["n_perm / seed", "9999 / 20260813", "9999 / 20260813"],
        ],
        widths=[2.3, 2.1, 2.1],
        caption="Table 4.8. P0.1 (p01_report.md; summary.csv / summary.json).",
    )
    r.h3_block("f", "Decision")
    r.p(
        "NULL after residualization. The observed cosine is on the wrong side of the null mean "
        "and the CI includes 0. LOSO passing does not override a non-significant, sub-null alignment. "
        "Phases 0–6 were not modified."
    )
    r.h3_block("g", "Output artifacts")
    r.bullets(
        [
            "results/similarity/p01_deviation/p01_report.md",
            "results/similarity/p01_deviation/summary.csv, summary.json",
            "cosine_matrix.csv, subject_alignment.csv, permutation_null.csv, figures/",
            "src/gait_research/similarity/deviation.py; scripts/run_p01_deviation.py",
        ]
    )

    r.h("4.9 P0.2 — Shared abnormality-set overlap", 2)
    r.h3_block("a", "Question")
    r.p(
        "Do the 17 victims share which preregistered features fall outside the control 10th–90th "
        "percentile band — a discrete abnormality set — even if continuous deviation directions "
        "(P0.1) do not align?"
    )
    r.h3_block("b", "Input")
    r.p(
        "Thirty locked features from Phase 2 subject_features.parquet (*__median / coordination), "
        "preferring Phase 3 redundancy representatives when available. Residualization uses the "
        "same four covariates, applied to continuous features before control-band binarization."
    )
    r.h3_block("c", "Method")
    r.p(
        "On the 14 controls only, compute the 10th–90th percentile band per feature. Each subject "
        "becomes a 30-bit exceedance vector (outside the band). Primary statistic: mean pairwise "
        "Jaccard among the 17 victims (greater than chance). Null: 9999 subject-label permutations, "
        "seed 20260813. Per-feature co-exceedance among victims plus BH-FDR across the 30 features "
        "is a diagnostic fingerprint, not a second primary claim. LOSO evaluates top-5 feature-rank "
        "agreement. Gate: post-residual perm p ≤ 0.05, LOSO feature-rank pass, and observed Jaccard "
        "above the permutation null mean."
    )
    r.h3_block("d", "Pre-registration")
    r.p(
        "Feature list locked in results/similarity/p02_abnormality/preregistered_features.json "
        "before any real test (locked_date 2026-08-16, n_features = 30, "
        "preregistered_before_any_real_test true). Selection rule: ax1 ROM bilateral hip/knee/ankle/"
        "foot-progress (8), ax1 peak timing (6), ax1 min timing (2), stance/cycle/mid-stance temporal (3), "
        "COM/path3d (3), and 4 ipsilateral coordination pairs × corr+lag (8). No post-hoc search of "
        "the 3665-column matrix."
    )
    r.h3_block("e", "Result")
    r.p(
        "Post-residual mean pairwise Jaccard 0.19059771741886694 (report 0.1906; synthesis 0.191) "
        "versus null mean 0.20118328610835212, permutation p = 0.6032. Observed is below the null "
        "mean. Features with co-exceedance FDR q ≤ 0.10: 0. LOSO top-5 feature-rank agreement 0.9677 "
        "but LOSO pass False (top-5 unstable by the test’s pass rule). Mean victim exceedance "
        "prevalence 0.292. Pre-residual Jaccard 0.1916 versus null 0.1889, p = 0.449, also not a gate pass."
    )
    r.table(
        ["Metric", "Pre-residual", "Post-residual (primary)"],
        [
            ["Mean pairwise Jaccard", "0.19156885473613983", "0.19059771741886694"],
            ["95% bootstrap CI", "[0.1920, 0.3054]", "[0.2013, 0.2925]"],
            ["Permutation p (greater)", "0.449", "0.6032"],
            ["Null mean / 95th pct", "0.1889 / 0.2466", "0.2012 / 0.2595"],
            ["Mean victim exceedance prevalence", "0.282", "0.292"],
            ["LOSO top-5 rank agree / pass", "0.742 / False", "0.968 / False"],
            ["Co-exceedance FDR q ≤ 0.10", "0 / 30", "0 / 30"],
            ["n_perm", "9999", "9999"],
        ],
        widths=[2.4, 2.05, 2.05],
        caption="Table 4.9. P0.2 (p02_report.md; summary.csv). CI values as in the primary report table.",
    )
    r.h3_block("f", "Decision")
    r.p(
        "NULL after residualization. Victims do not share which locked features exceed the control "
        "band more than chance. Phases 0–6 were not modified."
    )
    r.h3_block("g", "Output artifacts")
    r.bullets(
        [
            "results/similarity/p02_abnormality/p02_report.md",
            "preregistered_features.json, summary.csv, exceedance_matrix.csv, feature_coexceedance_residualized.csv",
            "src/gait_research/similarity/abnormality.py; scripts/run_p02_abnormality.py",
        ]
    )

    r.h("4.10 P0.3 — Amplitude-normalized waveform shape", 2)
    r.h3_block("a", "Question")
    r.p(
        "Do the 17 victims share waveform shape / timing on core gait curves after discarding "
        "ROM/amplitude — a question neither P0.1 (PC-space direction) nor P0.2 (binary exceedance) tested?"
    )
    r.h3_block("b", "Input")
    r.p(
        "Phase 1 frozen normalized_core.npz, summarized as subject nanmedian over cycles "
        "(label-blind), not recomputed from the raw MAT. Twelve locked curves. Residualization "
        "is applied to each (curve × phase %) column across subjects before z-scoring."
    )
    r.h3_block("c", "Method")
    r.p(
        "Z-score each subject-median curve across the 101 phase points (zero mean, unit variance), "
        "independently per curve. This discards amplitude/ROM and DC offset; only shape/timing remains. "
        "The same z-scored curves enter both Pearson and DTW so DTW cannot re-introduce magnitude. "
        "Two primary statistics, never averaged: (1) mean over curves of mean pairwise Pearson among "
        "victims (greater); (2) mean over curves of mean pairwise DTW distance (less). Null: 9999 "
        "subject-label permutations. BH-FDR across the 12 curves per measure. Pearson also requires "
        "LOSO pass for a defensible claim. Residualizing cycle duration removes linear associations "
        "between absolute gait speed and the value of each phase-% sample; DTW on 0–100% curves already "
        "allows nonlinear phase warping. These are different timing constructs; both are reported "
        "after the same residualization."
    )
    r.h3_block("d", "Pre-registration")
    r.p(
        "Curve list locked in results/similarity/p03_shape/preregistered_curves.json (locked_date "
        "2026-08-16, n_curves = 12). Curves: L/R Hip, Knee, Ankle, FootProgress ax1; LASI_ax1 and "
        "RASI_ax1 as Phase-1 pelvis proxies (LPelvisAngles/RPelvisAngles are not in the core cube); "
        "CentreOfMass ax1 and ax3. No post-hoc curve search. FDR family = 12 preregistered curves."
    )
    r.h3_block("e", "Result")
    r.p(
        "Pre-residual Pearson 0.5524147535678829 versus null 0.5533161067105464, p = 0.4936 — high "
        "absolute shape similarity that matches the null (generic gait shape). Post-residual Pearson "
        "−0.023369891483239296 versus null −0.029603465500581853, p = 0.2669; LOSO sign agreement "
        "0.8709677419354839, pass False. Post-residual DTW 7.656997577764042 versus null 7.590725909326609, "
        "p = 0.6288 (observed slightly worse than the null). Curves with FDR q ≤ 0.10: 0 and 0. "
        "Synthesis rounds Pearson to −0.023 / DTW 7.657, p 0.267 / 0.629."
    )
    r.table(
        ["Metric", "Pearson pre", "Pearson post", "DTW pre", "DTW post"],
        [
            ["Mean pairwise", "0.5524", "−0.0234", "4.0376", "7.6570"],
            ["95% CI", "[0.5099, 0.7005]", "[−0.0439, 0.2584]", "[3.1591, 4.2377]", "[5.9529, 7.7332]"],
            ["Permutation p", "0.4936 (>)", "0.2669 (>)", "0.4902 (<)", "0.6288 (<)"],
            ["Null mean / tail", "0.5533 / p95=0.5973", "−0.0296 / p95=0.0178", "4.0312 / p05=3.7599", "7.5907 / p05=7.3056"],
            ["LOSO pass", "True", "False", "—", "—"],
            ["FDR q ≤ 0.10", "0/12", "0/12", "0/12", "0/12"],
        ],
        widths=[1.3, 1.3, 1.4, 1.25, 1.25],
        caption="Table 4.10. P0.3 (p03_report.md; summary.csv). Rounded as in the primary report table; CSV has full precision.",
    )
    r.h3_block("f", "Decision")
    r.p(
        "NULL after residualization. Gate requires at least one of Pearson or DTW with perm p ≤ 0.05 "
        "and observed on the similarity side of the null, with Pearson also requiring LOSO pass. "
        "Neither measure meets it. Pre-residual high Pearson is not victim-specific."
    )
    r.h3_block("g", "Output artifacts")
    r.bullets(
        [
            "results/similarity/p03_shape/p03_report.md, preregistered_curves.json, summary.csv",
            "per_curve_similarity_residualized.csv, permutation_null.csv",
            "src/gait_research/similarity/shape_space.py; scripts/run_p03_shape.py",
        ]
    )

    r.h("4.11 P0.4 — Event-localized phase windows", 2)
    r.h3_block("a", "Question")
    r.p(
        "Is victim similarity localized to specific clinical gait phases that whole-cycle tests "
        "(P0.1–P0.3) dilute into noise?"
    )
    r.h3_block("b", "Input")
    r.p(
        "Phase 1 gait_cycle_inventory.csv events on all 880 cycles, and the 12 P0.3 locked curves "
        "from normalized_core.npz. Windowing reuses deviation.py (cosine) and abnormality.py (Jaccard)."
    )
    r.h3_block("c", "Method")
    r.p(
        "Five event-bounded windows, not fixed 10% bins: loading response (IC → opposite FO), "
        "mid-stance (opposite FO → Midsvnt), terminal stance (Midsvnt → opposite FC), pre-swing "
        "(opposite FC → ipsilateral FO), undivided swing (ipsilateral FO → next IC). For each "
        "window × curve, subject-level mean and ROM are computed, then tested with deviation cosine "
        "and abnormality Jaccard. FDR family, stated before running: 240 = 5 phases × 12 curves × "
        "2 aggregations × 2 tests; BH-FDR spans the entire family, not per-window. Residualize "
        "continuous window features on the four covariates before the primary gate. Secondary: "
        "per-window 24-D multivariate (12 curves × mean/ROM) cosine and Jaccard with LOSO. "
        "Permutations: 9999, seed 20260813. Gate: ≥1 cell with FDR q ≤ 0.10 after residualization "
        "and supporting window-level LOSO for that phase."
    )
    r.h3_block("d", "Pre-registration")
    r.p(
        "Locked in results/similarity/p04_event_phases/preregistered_phases.json (locked_date "
        "2026-08-16). Audit: 100% of 880 cycles have all listed events in strict temporal order. "
        "Explicitly not reconstructable and not estimated: initial / mid / terminal swing (no "
        "feet-adjacent, tibia-vertical, or equivalent events). Decision in the lock file: use the "
        "5 reconstructable windows only; do not interpolate ISw/MSw/TSw. n_fdr_family = 240."
    )
    r.h3_block("e", "Result")
    r.p(
        "Post-residual: FDR family size 240; cells with q ≤ 0.10: 0; cells with q ≤ 0.05: 0; "
        "minimum raw permutation p = 0.0188 (summary.csv; report 0.0188; synthesis rounds to 0.019). "
        "A minimum raw p of 0.0188 is expected under 240 comparisons and does not survive FDR. "
        "Pre-residual: 0 at q ≤ 0.10, min raw p = 0.0279. Per-window multivariate residualized "
        "cosines are all negative (less aligned than a typical null), all cosine p ≥ 0.8476; "
        "Jaccard p ≥ 0.5443. Swing cosine −0.0523, p = 0.968."
    )
    r.table(
        ["Metric", "Pre-residual", "Post-residual (primary)"],
        [
            ["FDR family size", "240", "240"],
            ["Cells q ≤ 0.10", "0", "0"],
            ["Cells q ≤ 0.05", "0", "0"],
            ["Min raw perm p", "0.0279", "0.0188"],
            ["n_perm", "9999", "9999"],
        ],
        widths=[2.4, 2.05, 2.05],
        caption="Table 4.11a. P0.4 family-wide FDR (p04_report.md; summary.csv).",
    )
    r.table(
        ["Phase", "Cosine", "Cos p", "Cos LOSO", "Jaccard", "Jac p", "Jac LOSO"],
        [
            ["loading_response", "−0.0369", "0.9297", "False", "0.1474", "0.8805", "False"],
            ["mid_stance", "−0.0491", "0.9585", "True", "0.1494", "0.7795", "False"],
            ["terminal_stance", "−0.0588", "0.9987", "True", "0.1549", "0.8945", "False"],
            ["pre_swing", "−0.0395", "0.8476", "True", "0.1870", "0.5443", "False"],
            ["swing (undivided)", "−0.0523", "0.968", "True", "0.1526", "0.7953", "False"],
        ],
        widths=[1.4, 0.85, 0.8, 0.85, 0.85, 0.8, 0.85],
        caption="Table 4.11b. Post-residual per-window 24-D multivariate (p04_report.md).",
    )
    r.h3_block("f", "Decision")
    r.p(
        "NULL after residualization. Restricting to event-derived phases does not reveal a "
        "localized shared signature under the 240-cell FDR family."
    )
    r.h3_block("g", "Output artifacts")
    r.bullets(
        [
            "results/similarity/p04_event_phases/p04_report.md, preregistered_phases.json",
            "cell_results_residualized.csv, window_multivariate_residualized.csv, summary.csv",
            "src/gait_research/similarity/event_phases.py; scripts/run_p04_event_phases.py --n-perm 9999",
        ]
    )

    r.h("4.12 P0.5 — Confound residualization (folded in)", 2)
    r.h3_block("a", "Question")
    r.p(
        "Could a shared victim pattern be an artifact of body size or gait speed, and conversely, "
        "does residualizing those covariates create or destroy a similarity signal? P0.5 in the "
        "original plan was “confound as shared residual pattern.” In this project it is not a "
        "separate discovery test."
    )
    r.h3_block("b", "Input")
    r.p(
        "The same subject-level covariates in every P0 test that residualizes: height_cm, mass_kg, "
        "mean_leg_cm, cycle_duration_s_median. Continuous representations (Phase 4 27-D, P0.2 "
        "features, P0.3 curve×time columns, P0.4 window features, P0.6 wrapped CRP radians) are "
        "the variables residualized."
    )
    r.h3_block("c", "Method")
    r.p(
        "Subject-level OLS with intercept of each continuous column on the four covariates, "
        "across the 31 subjects, before the primary similarity statistic. Pre- and post-residual "
        "results are always reported; the decision gate uses post-residual. For P0.6, residualization "
        "of wrapped CRP radians is described as pragmatic confound control. P0.3 documents that "
        "residualizing duration is not a substitute for DTW warping."
    )
    r.h3_block("d", "Pre-registration")
    r.p(
        "The four-covariate list is the project-wide convention (README Similarity P0 table). "
        "There is no preregistered_p05.json and no scripts/run_p05_*.py."
    )
    r.h3_block("e", "Result")
    r.p(
        "There is no unique P0.5 test statistic. The effect of residualization is visible in each "
        "other test’s pre/post tables: P0.1 cosine 0.0640 → 0.0518 (both sub-null); P0.2 Jaccard "
        "essentially unchanged and still null; P0.3 Pearson 0.5524 → −0.0234 (generic shared shape "
        "removed with amplitude-related covariates; residual shape still null); P0.6 circular "
        "similarity 0.8374 → 0.6494 (still below the residualized null 0.6763)."
    )
    r.table(
        ["Item", "Value"],
        [
            ["Separate discovery script", "None"],
            ["Covariates", "height_cm, mass_kg, mean_leg_cm, cycle_duration_s_median"],
            ["Estimator", "OLS with intercept, subject level"],
            ["Gate", "Post-residual primary in P0.1–P0.4 and P0.6"],
            ["p0_synthesis.md gate column", "folded in"],
        ],
        widths=[2.6, 3.9],
        caption="Table 4.12. P0.5 status (p0_synthesis.md; each P0 report’s residualization block).",
    )
    r.h3_block("f", "Decision")
    r.p(
        "Folded in — not a discovery claim. Residualization is a confound control and a primary "
        "evidence gate, not a finding that victims share a residual pattern."
    )
    r.h3_block("g", "Output artifacts")
    r.bullets(
        [
            "No P0.5 results directory",
            "Pre/post tables in each of results/similarity/p01_deviation/ through p06_coordination/",
            "results/similarity/p0_synthesis.md row P0.5",
        ]
    )

    r.h("4.13 P0.6 — Continuous relative phase (CRP) coupling", 2)
    r.h3_block("a", "Question")
    r.p(
        "Do the 17 victims share inter-joint coupling (CRP profiles) that is invisible to "
        "single-curve / single-feature tests (P0.1–P0.4)?"
    )
    r.h3_block("b", "Input")
    r.p(
        "Phase 1 normalized_core.npz ax1 angle curves (101 points). Angular velocity is not stored "
        "in Phase 1 and is not required for Hilbert phase. Finite-difference / Savitzky–Golay "
        "velocity (used in Phase 2 features) is unused for Hilbert CRP."
    )
    r.h3_block("c", "Method")
    r.p(
        "Instantaneous phase is the Hilbert analytic signal of the demeaned angle "
        "(scipy.signal.hilbert → np.angle). Hilbert is preferred over atan2(ω, θ) because "
        "phase-plane methods need separate position/velocity normalization. CRP = wrap(φ_proximal "
        "− φ_distal) to (−π, π]. Subject profile = circular mean CRP across cycles (atan2 of mean "
        "sin/cos). Two statistics, FDR family 12 = 6 pairs × 2 measures: (1) circular mean_t "
        "cos(CRP_i − CRP_j), which preserves constant phase offsets that z-scored Pearson would "
        "destroy; (2) DTW on unwrap(CRP − CRP[0]) for time-varying coupling shape. Null: 9999 "
        "subject-label permutations, seed 20260813. Linear residualization of wrapped CRP radians "
        "on the four covariates. Gate: post-residual perm p ≤ 0.05 on circular similarity or DTW "
        "(circular/Pearson-path LOSO pass), observed on the similarity side of the null, and ≥1 "
        "pair×measure FDR q ≤ 0.10."
    )
    r.h3_block("d", "Pre-registration")
    r.p(
        "Pairs locked in results/similarity/p06_coordination/preregistered_pairs.json "
        "(locked_date 2026-08-17, n_pairs = 6): left/right hip–knee, knee–ankle, hip–ankle, all ax1. "
        "crp_method = hilbert_analytic_phase. n_fdr_family = 12."
    )
    r.h3_block("e", "Result")
    r.p(
        "Post-residual circular similarity 0.6493969168966108 (report 0.6494; synthesis 0.649) "
        "versus null mean 0.6762901470475436, permutation p = 0.929 — observed below the null. "
        "DTW 17.62505483572455 versus null 17.22483148462879, p = 0.6048 — observed slightly less "
        "similar than the null. Pairs with FDR q ≤ 0.10: 0 and 0. LOSO pass True, sign agreement 1.000 "
        "on the circular statistic. Pre-residual circular 0.8374 versus null 0.8395, p = 0.5536: high "
        "absolute coupling similarity is shared gait coordination, not victim-specific excess."
    )
    r.table(
        ["Metric", "Circular pre", "Circular post", "DTW pre", "DTW post"],
        [
            ["Mean pairwise", "0.8374", "0.6494", "16.6049", "17.6251"],
            ["95% CI", "[0.8086, 0.8853]", "[0.6241, 0.7233]", "[12.3493, 17.3494]", "[13.1688, 18.9195]"],
            ["Permutation p", "0.5536", "0.929", "0.6883", "0.6048"],
            ["Null mean", "0.8395", "0.6763", "16.0354", "17.2248"],
            ["LOSO pass / sign", "True / 1.000", "True / 1.000", "—", "—"],
            ["FDR q ≤ 0.10", "0/6", "0/6", "0/6", "0/6"],
        ],
        widths=[1.3, 1.3, 1.35, 1.25, 1.3],
        caption="Table 4.13. P0.6 (p06_report.md; summary.csv). Circular column in CSV is named mean_pairwise_pearson for the circular cos statistic.",
    )
    r.h3_block("f", "Decision")
    r.p(
        "NULL after residualization. High absolute circular similarity (~0.65–0.84) is normal "
        "gait coordination when it does not exceed the subject-label null. P1 was not started."
    )
    r.h3_block("g", "Output artifacts")
    r.bullets(
        [
            "results/similarity/p06_coordination/p06_report.md, preregistered_pairs.json, summary.csv",
            "per_pair_similarity_residualized.csv",
            "src/gait_research/similarity/coordination_crp.py; scripts/run_p06_coordination.py",
        ]
    )

    r.h("4.14 Power / MDE analysis (P0.1 only)", 2)
    r.h3_block("a", "Question")
    r.p(
        "At n = 17 victims versus 14 controls in residualized Phase 4 27-D space, how large a "
        "shared deviation direction would P0.1 have been likely to detect? This does not re-run "
        "or alter the frozen P0.1 test. It asks whether the observed null (cosine 0.0518, p = 0.7579) "
        "is unsurprising given the design’s sensitivity."
    )
    r.h3_block("b", "Input")
    r.p(
        "Frozen residualized 31 × 27 Phase 4 cloud (same residualization as P0.1 primary). Frozen "
        "P0.1 residualized permutation null for the shape check (recomputed in the power folder as "
        "residualized_p01_perm_null.csv, not by rewriting p01_deviation/). Typical individual scale: "
        "median control ||d_i|| = 8.382963989053888 (report 8.3830)."
    )
    r.h3_block("c", "Method")
    r.p(
        "Effect size λ is shared-offset Euclidean length as a fraction of median control ||d||: "
        "offset = λ × median_control||d|| along a random unit direction (same generative pattern as "
        "tests/similarity/test_deviation.py::test_shared_direction_detected). Grid: 0, 0.25, 0.5, "
        "0.75, 1.0, 1.25, 1.5, 2.0, 3.0. At each λ, 1000 datasets are simulated: randomly partition "
        "the empirical 31 points into 17/14 and, for λ > 0, add the shared offset to the 17. Each "
        "dataset is tested with the actual P0.1 statistic and permutation test (999 permutations, "
        "documented reduction from P0.1’s 9999). Power is the proportion of simulations with "
        "permutation p ≤ 0.05. MDE is the interpolated λ where power crosses 0.80. Seed 20260813. "
        "A parametric Ledoit–Wolf MVN fit to the 14 controls failed the null-shape check (too "
        "spherical; simulated λ = 0 cosine mean ~0.06 versus residualized permutation null 0.11) "
        "and was not used for the headline MDE. At λ = 0, empirical injection is the P0.1 "
        "permutation experiment, so the cosine distribution must match by construction."
    )
    r.h3_block("d", "Pre-registration")
    r.p(
        "The MDE analysis is post-hoc relative to P0.1 (it does not alter the frozen test). "
        "α = 0.05, 80% power convention, n_sim = 1000, n_perm = 999, and the λ grid are documented "
        "in mde_report.md and power_analysis.py. Power was not estimated for P0.2–P0.6."
    )
    r.h3_block("e", "Result")
    r.p(
        "False-positive rate at λ = 0: 0.052 (52/1000; target ~0.05). Power at λ = 3.0: 1.000. "
        "Simulated λ = 0 cosine mean 0.10911835114146004 versus real residualized P0.1 null mean "
        "0.11029207303487745 (relative difference 0.0106); SDs 0.07627 versus 0.07713 (relative "
        "difference 0.0112). Headline MDE: 80% power at λ = 0.7284560143626572 (report 0.73), "
        "with interpolated mean observed cosine 0.2996658202861148 (report 0.300). Observed P0.1 "
        "cosine 0.05181913119541501 is smaller than that detection threshold. At λ = 0.50, power "
        "was only 0.291. Victims would need a common offset of 0.73× a typical control’s entire "
        "residual deviation before this n = 17 vs 14 design reliably detects it."
    )
    r.table(
        ["λ", "Power", "Mean observed cosine", "Rejections"],
        [
            ["0.00", "0.052", "0.10911835114146004", "52/1000"],
            ["0.25", "0.091", "0.140430385916014", "91/1000"],
            ["0.50", "0.291", "0.21133519562055295", "291/1000"],
            ["0.75", "0.848", "0.3079956237909222", "848/1000"],
            ["1.00", "0.995", "0.3994381814397216", "995/1000"],
            ["1.25", "1.000", "0.47518138706253404", "1000/1000"],
            ["1.50", "1.000", "0.5471607957575649", "1000/1000"],
            ["2.00", "1.000", "0.6533676181646951", "1000/1000"],
            ["3.00", "1.000", "0.7889606377792873", "1000/1000"],
            ["MDE (interp.)", "0.80", "0.2996658202861148", "λ = 0.7284560143626572"],
        ],
        widths=[1.4, 1.2, 2.3, 1.6],
        caption="Table 4.14. P0.1 power curve (power_analysis/summary.csv and summary.json).",
    )
    r.h3_block("f", "Decision")
    r.p(
        "The MDE is a limitation statement, not a positive P0.1 finding. It does not turn the "
        "P0.1 null into a positive. λ > 0 is a shared Euclidean offset in a random direction on "
        "the empirical residual cloud — one alternative, not every possible shared pattern. "
        "Phases 0–6 and P0.1–P0.6 result files were not modified."
    )
    r.h3_block("g", "Output artifacts")
    r.bullets(
        [
            "results/similarity/power_analysis/mde_report.md",
            "summary.csv, summary.json, figures/power_curve.png, figures/null_compare.png",
            "src/gait_research/similarity/power_analysis.py; scripts/run_power_analysis.py",
            "tests/similarity/test_power_analysis.py (calibration assertions)",
        ]
    )

    # ============================================================
    # 5. Cross-cutting
    # ============================================================
    r.h("5. Cross-Cutting Design Rules", 1)
    r.h("Subject as the independent unit", 2)
    r.p(
        "Pseudo-replication is the central methodological risk in this archive. 880 cycles and "
        "260 trials are repeated measures within 31 people. If cycles were shuffled or modeled as "
        "independent, every p-value in Phases 3–6 and P0 would be anti-conservative. The rule is "
        "therefore operational, not rhetorical: quality screens, permutations, clustering stability "
        "subsamples, LOSO, and bootstrap CIs all operate on subjects. Phase 6 keeps each person’s "
        "101-point trajectory intact and shuffles labels, never time. Certification checks named "
        "no_cycle_as_subject, perm_unit_subject, not_880, and analogous flags exist in P0 reports "
        "(“Labels shuffled across subjects only”)."
    )
    r.h("Label-blind preprocessing, and where labels first enter", 2)
    r.p(
        "Phase 0 uses labels only to verify the 17/14 split and join. Phase 1 extraction does not "
        "use labels to accept or reject cycles (usable_lower_body is a coverage/duration rule). "
        "Phase 2 forbids victimization tokens in feature code and tables. Phase 3 screening and "
        "redundancy raise an error if a victimized column is present; labels join only for "
        "Mann–Whitney / δ / permutation. Phase 4 omits labels from scaling, PCA, k selection, "
        "stability, and characterization; labels join after assignments freeze. Phase 5 uses labels "
        "only to subset the 17 and to test similarity. Phase 6 uses labels only in the permutation "
        "test after subject medians exist. P0 uses labels to compute victim-subset similarity and "
        "to permute those labels; curve/feature/phase/pair lists are locked before group results. "
        "A classifier was never used as a discovery tool in any layer."
    )
    r.h("Confound residualization", 2)
    r.p(
        "The four covariates are height_cm, mass_kg, mean_leg_cm, and cycle_duration_s_median. "
        "They are the anthropometric and tempo quantities most likely to move whole kinematic "
        "vectors without being victimization itself. Residualization is subject-level OLS with "
        "intercept, applied before the primary P0 statistic. Phase 4 separately reports Kruskal–Wallis "
        "of anthropometry versus the unsupervised phenotype (height p = 0.0499). Residualization "
        "is not a claim that these four exhaust confounding (lab protocol, footwear, unmarked "
        "axes, survey heterogeneity remain)."
    )
    r.h("Multiplicity control, by family", 2)
    r.table(
        ["Family", "Size", "Method", "Survivors at the stated gate"],
        [
            ["Phase 3 representative tests", "335", "BH-FDR on Mann–Whitney raw p", "0 at q ≤ 0.10 and q ≤ 0.05"],
            ["Phase 4 phenotype enrichment", "2 phenotypes", "FDR on subject-label perm p", "both q = 1"],
            ["Phase 6 time points within channel", "101", "cluster permutation (max mass)", "embedded in region class"],
            ["Phase 6 signals within analysis level", "primary / secondary / asymmetry", "BH-FDR", "0 ROBUST / 0 EXPLORATORY"],
            ["Phase 6 shape descriptors", "shape battery", "Mann–Whitney + BH", "none FDR ≤ 0.10"],
            ["P0.2 co-exceedance", "30 locked features", "BH-FDR", "0 at q ≤ 0.10"],
            ["P0.3 per-curve Pearson / DTW", "12 + 12", "BH-FDR", "0 at q ≤ 0.10"],
            ["P0.4 window cells", "240", "BH-FDR on the entire family", "0 at q ≤ 0.10"],
            ["P0.6 pair × measure", "12", "BH-FDR", "0 at q ≤ 0.10"],
            ["P0.1 primary cosine", "1", "subject-label permutation (no FDR family)", "p = 0.7579, NULL"],
        ],
        widths=[2.2, 1.3, 1.6, 1.4],
        caption="Table 5.1. Multiplicity families actually used.",
    )
    r.h("Vocabulary: robust, exploratory, certified, null", 2)
    r.p(
        "These words are project-specific and should not be read as generic English."
    )
    r.bullets(
        [
            "Certified / PASS: the pipeline ran the specified methods (unit, label timing, files present). It does not mean a scientific positive. Phase 3–6 certifications are PASS WITH WARNINGS when the honest result is no signature.",
            "Signature (Phase 3): a feature that meets all of FDR q ≤ 0.10, |δ| ≥ 0.33, LOSO direction ≥ 0.80, victim consistency ≥ 0.60. Zero features met this.",
            "Exploratory (Phase 3): a ranked candidate that failed the signature rule. It is a ranking aid, not a confirmed victim gait.",
            "ROBUST (Phase 6): a trajectory region meeting the full predefined robustness list (primary channel, cluster perm + FDR, medium+ effect, consistency, LOSO, CI excluding 0, contiguous span). Count = 0.",
            "EXPLORATORY (Phase 6): a weaker pre-specified trajectory class (the report records 0). Not the same as Phase 3 “exploratory ranks.”",
            "UNSUPPORTED (Phase 6): a cluster worth listing (e.g. lowest p) that does not meet ROBUST/EXPLORATORY. Not a finding.",
            "Outgroup (Phase 4): the n = 4 side of a 27 vs 4 split. Explicitly not a named clinical phenotype.",
            "NULL (P0): the pre-registered discovery gate failed after residualization (and LOSO/FDR where applicable). High absolute similarity that matches the permutation null is still NULL.",
            "Folded in (P0.5): residualization was applied inside other tests; there is no separate P0.5 p-value.",
            "MDE / 80% power: a post-hoc sensitivity number for P0.1, not a new similarity discovery.",
        ]
    )

    # ============================================================
    # 6. Synthesis
    # ============================================================
    r.h("6. Synthesis Across All Steps", 1)
    r.p(
        "The two research questions are answered separately. Merging them into a single slogan "
        "(“gait is unrelated to victimization”) would over-claim: this sample, these constructs, "
        "and this power are what was tested."
    )
    r.table(
        ["ID", "Construct", "Primary post-residual statistic", "Null mean", "Perm p", "FDR", "LOSO", "Gate"],
        [
            ["P0.1", "Deviation direction (27-D)", "cosine 0.0518", "0.1103", "0.7579", "—", "pass", "NULL"],
            ["P0.2", "Abnormality set (30 features)", "Jaccard 0.1906", "0.2012", "0.6032", "0/30", "top-5 fail", "NULL"],
            ["P0.3", "Waveform shape (12 curves)", "Pearson −0.0234 / DTW 7.6570", "−0.0296 / 7.5907", "0.2669 / 0.6288", "0/12", "fail", "NULL"],
            ["P0.4", "Event-phase windows", "240-cell FDR family", "—", "min raw 0.0188", "0 at q≤0.10", "window MV null", "NULL"],
            ["P0.5", "Confound residualization", "not a separate test", "—", "—", "—", "—", "folded in"],
            ["P0.6", "CRP coupling (6 pairs)", "circ. 0.6494 / DTW 17.6251", "0.6763 / 17.2248", "0.929 / 0.6048", "0/12", "circ. pass", "NULL"],
        ],
        widths=[0.7, 1.35, 1.35, 0.85, 0.85, 0.7, 0.85, 0.7],
        caption="Table 6.1. Similarity P0 decision-gate summary (p0_synthesis.md; primary-file precision).",
    )
    r.table(
        ["Phase", "Question type", "Headline numeric result", "Certification", "Scientific call"],
        [
            ["0", "Audit", "31 subjects; 17/14; 260 WU*; 0 critical / 74 warnings", "PASS WITH WARNINGS", "Data usable"],
            ["1", "Extraction", "880/880 usable cycles; KinFC 1=R, 2=L", "PASS", "Cycles exist"],
            ["2", "Features", "714 cycle / 3665 subject columns; 0 label leak", "PASS", "Representation frozen"],
            ["3", "Mean differences", "0/335 FDR q≤0.10; 0 signature-rule hits", "PASS WITH WARNINGS", "No shared mean signature"],
            ["4", "Phenotypes", "k=2, 27 vs 4; enrichment p=1; height Kruskal p=0.0499", "PASS WITH WARNINGS", "Outgroup, not a victim phenotype"],
            ["5", "Victim compactness", "distance 20.417 vs null 20.060; p=0.547; 0 subgroups", "PASS WITH WARNINGS", "Victims not a neighborhood"],
            ["6", "Trajectories", "0 ROBUST; 0 EXPLORATORY; strongest p=0.1469", "PASS WITH WARNINGS", "No time-localized difference"],
        ],
        widths=[0.7, 1.2, 2.3, 1.2, 1.1],
        caption="Table 6.2. Parallel Phases 0–6 decisions (primary reports and certifications).",
    )
    r.h("What effect size this program could and could not have detected", 2)
    r.p(
        "The quantified sensitivity statement is for P0.1, the most interpretable shared-direction "
        "test. 80% power required λ ≥ 0.7285 ≈ 0.73× typical control residual ||d||, corresponding "
        "to expected cosine ≈ 0.30. The observed cosine 0.0518 is far below that bar. At λ = 0.50 "
        "power was 29%. Power was not simulated for P0.2–P0.6; those nulls are still valid tests "
        "of their constructs, but they do not carry a matching MDE number. Phase 3 already warned "
        "that n = 17 vs 14 has low power and that a true medium Cliff’s δ can fail FDR among 335 tests. "
        "Taken together: the program was capable of detecting a gross shared deviation direction "
        "and did not; it was not powered to rule out subtle shared patterns or modest mean differences "
        "after heavy multiplicity."
    )
    r.h("Two separate conclusions (not merged)", 2)
    r.p(
        "Q1 — individual / group-difference signature (Phases 0–6): In this sample, no biomechanical "
        "gait characteristic showed a statistically supported, robust, interpretable, and "
        "subject-consistent difference between victimized and non-victimized females, as those "
        "words are defined by the signature rule, enrichment tests, compactness tests, and "
        "trajectory robustness classes actually used."
    )
    r.p(
        "Q2 — shared-victim-pattern signature (Similarity P0): In this sample, the 17 victimized "
        "women did not share a locomotor pattern with each other that controls do not, on the "
        "pre-registered constructs of deviation direction, abnormality set, amplitude-normalized "
        "shape, event-phase windows, or Hilbert CRP coupling, after residualization, LOSO where "
        "defined, and BH-FDR where a family was locked."
    )
    r.p(
        "These conclusions are about this cohort, these methods, and this power. They are not a "
        "license to start P1 on the same 31 subjects to hunt a positive with more elaborate metrics. "
        "p0_synthesis.md: given six consecutive null decision gates (P0.1–P0.4, P0.6; P0.5 absorbed), "
        "starting P1 on the same cohort would be exploratory dredging unless there is a new, "
        "independently motivated hypothesis and preferably an external cohort. Recommendation: "
        "do not start P1 without explicit go-ahead."
    )

    # ============================================================
    # 7. Limitations
    # ============================================================
    r.h("7. Limitations", 1)
    r.p(
        "README §10 lists nine limitations. Each is expanded here with the reasoning that makes "
        "it a constraint on citation, not a footnote."
    )
    r.h("7.1 Sample size and the quantified MDE", 2)
    r.p(
        "n = 31 (17 vs 14) is small for high-dimensional gait. Phase 3 said so qualitatively "
        "(a true medium effect can fail FDR among 335 tests). The P0.1 simulation made that "
        "quantitative: 80% power only for a shared residual deviation of magnitude ≥ 0.73× a "
        "typical control’s ||d_i|| (expected cosine ≈ 0.30). Observed cosine 0.0518 is well below "
        "that MDE, so the P0.1 null is unsurprising given the design. Absence of evidence is not "
        "proof that no gait difference exists in a larger population. It is proof that this design "
        "did not detect a gross shared direction and was unlikely to detect a subtle one."
    )
    r.h("7.2 Survey-label heterogeneity", 2)
    r.p(
        "Victimization is a survey field (Y/N), not a randomized or clinician-adjudicated exposure. "
        "Among Y = 17, VictimType is Nd 7, Ip 6, Both 3, online 1; Times is sometimes Nd even when "
        "Y (S19, S26). In-person assault, online-only experience, both, and undisclosed type are "
        "not the same construct. Pooling them as one “victimized” class can dilute a subtype-specific "
        "motor pattern if one exists, and can also mix unrelated life histories that have nothing "
        "to do with gait. The project did not run subtype-stratified discovery tests as confirmatory "
        "inference; n per cell is too small (online-only n = 1)."
    )
    r.h("7.3 Laboratory gait versus daily walking", 2)
    r.p(
        "Capture is a Plug-in Gait protocol at 100 Hz, overground walking trials labeled WU*, "
        "not free-living locomotion. Laboratory gait is slower, more self-conscious, and more "
        "constrained in path and footwear than daily walking. A null in the lab does not prove "
        "a null in the street; a lab difference would still have been a lab difference. This "
        "program only speaks to the recorded protocol."
    )
    r.h("7.4 S15 stature and low cycle counts", 2)
    r.p(
        "S15 is 141.0 cm (blueprint), a short outlier relative to the 141.0–167.6 cm range, and "
        "is victimized (Both). Subject medians for S15 can pull spatial and some angular features. "
        "Residualizing height is a partial control, not a guarantee. S19 (victim, Ip) has 4 usable "
        "cycles from 4 trials; S30 (control) has 3 usable cycles from 5 trials. Their subject "
        "medians are noisier than a person with 40 cycles. Phase 4’s n = 4 outgroup includes S19 "
        "and S5 (both Y) plus S35 and S40 (both N), which is another reason not to name that split "
        "a victim phenotype."
    )
    r.h("7.5 Unnamed coordinate axes", 2)
    r.p(
        "Markers and angles are stored as ax1/ax2/ax3. Phase 2 certification explicitly requires "
        "that feature names not say anterior–posterior / mediolateral / vertical. Interpreting a "
        "“sagittal hip ROM” finding would require a separate coordinate-convention study. P0.3/P0.6 "
        "use ax1 as the primary angle axis by lock-file convention, not because ax1 is certified sagittal."
    )
    r.h("7.6 Phase 4 outgroup mislabeling risk", 2)
    r.p(
        "k = 2 met stability (silhouette 0.415, mean bootstrap ARI 0.755, min size 4) but produced "
        "27 vs 4. The source report forbids promoting the n = 4 group to a named clinical phenotype. "
        "Doing so would convert a majority/outgroup split — possibly stature-related (height Kruskal "
        "p = 0.0499) — into a false “victim gait type.” Enrichment p-values are 1."
    )
    r.h("7.7 Non-independence of any future classifier on this cohort", 2)
    r.p(
        "Every feature, cluster, curve, pair, and phase list in this repository was examined in "
        "the service of Q1 or Q2 on these 31 people. A supervised model trained and tested on the "
        "same 31, even with cross-validation, would not be an independent confirmation of anything "
        "discovered here. The honest next classifier lives on new subjects, with a frozen "
        "pre-registered feature list."
    )
    r.h("7.8 High absolute CRP / Pearson similarity is generic gait", 2)
    r.p(
        "Humans walking overground share hip–knee–ankle coupling and waveform shape. P0.3 "
        "pre-residual Pearson 0.5524 matched null 0.5533; P0.6 pre-residual circular similarity "
        "0.8374 matched null 0.8395. Those numbers look like “victims are similar” until the "
        "permutation null is shown. Similarity P0 therefore reports observed versus null, not "
        "absolute similarity alone. A reader who cites 0.65 circular CRP without the p = 0.929 "
        "would be mis-citing the project."
    )
    r.h("7.9 Pelvis proxy and unreconstructed swing phases", 2)
    r.p(
        "Phase 1 core has pelvis markers, not PelvisAngles. P0.3 and P0.4 used LASI/RASI ax1 as "
        "documented proxies in the lock files. Initial, mid, and terminal swing were not interpolated "
        "for P0.4 because stored events cannot split swing. A phase-local effect confined to mid-swing "
        "would have been invisible by construction. That is a coverage gap, not a negative test of those phases."
    )

    # ============================================================
    # 8. What this is not
    # ============================================================
    r.h("8. What This Project Is Not", 1)
    r.p("README §11 is reproduced and expanded with one sentence of reasoning per item.")
    r.p(
        "Not a diagnostic tool. No threshold, score, or clinical cut-point was derived, and the "
        "inferential results are null; using these files to “flag” victimization from gait would "
        "be a use the evidence does not support."
    )
    r.p(
        "Not evidence of causation. Victimization labels are survey-reported, temporally unordered "
        "relative to the gait session in the analyses, and unconfounded only to the limited extent "
        "of four residualized covariates. Even a positive would have been an association."
    )
    r.p(
        "Not a claim that “gait identifies victims.” Identification is a classification task that "
        "was deliberately not trained. The shared-pattern tests that could have supported a "
        "weaker version of that claim were null."
    )
    r.p(
        "Not a license to pool 880 cycles as n. Doing so would manufacture significance by "
        "pseudo-replication. Every certified permutation unit in this project is the subject."
    )
    r.p(
        "Not Phase 7 (supervised prediction). No XGBoost, neural net, or victim score has been "
        "trained. Predictive modeling was deferred because n = 31 independent people."
    )
    r.p(
        "Not Similarity P1 (Wasserstein / RV / soft-DTW / common subspace). Those methods were "
        "explicitly not started after P0 nulls, because escalating metric complexity on the same "
        "already-tested sample is dredging unless a new hypothesis and preferably a new cohort exist."
    )

    # ============================================================
    # 9. Future work
    # ============================================================
    r.h("9. Recommended Future Work", 1)
    r.h("New or larger cohort", 2)
    r.p(
        "The binding constraint is independent people, not more cycles per person. A confirmatory "
        "study should pre-register a short feature/curve/pair list before seeing labels, use the "
        "subject as the unit, residualize (or match) body size and tempo, and size the sample to "
        "the MDE it cares about. If the scientific target is a subtle shared direction (λ ≪ 0.73), "
        "n = 17 vs 14 is the wrong design; the P0.1 curve is the planning input."
    )
    r.h("Other modalities", 2)
    r.p(
        "This archive is Plug-in Gait kinematics. EMG, force plates, and upper-body markers with "
        "adequate coverage (Phase 0 documented LUPA/RUPA gaps) could test different mechanisms "
        "(muscle timing, kinetics, arm carriage) that joint-angle kinematics may not show. They "
        "would still need subject-level inference and pre-registration."
    )
    r.h("Subtype-stratified description, not a fishing expedition", 2)
    r.p(
        "VictimType (Ip / online / Both / Nd) is real heterogeneity. With n = 1 online-only and "
        "n = 3 Both, confirmatory subtype tests on this file are not powered. Descriptive tables "
        "of subtype × anthropometry × cycle count can be groundwork for a future study that "
        "recruits to those cells. They should not be used to hunt a significant subtype after "
        "the binary tests were null, unless that hunt is labeled exploratory and not cited as a finding."
    )
    r.h("P1 on this cohort is not recommended", 2)
    r.p(
        "Wasserstein distances on distributions of cycles, RV coefficients on covariance, soft-DTW, "
        "and common-subspace methods are legitimate constructs. They are also additional families "
        "of tests on the same 31 labels already used in Phase 3 (335 tests), Phase 4–6, and six P0 "
        "gates. Starting P1 because P0 was null is the definition of escalating complexity until "
        "a positive appears. p0_synthesis.md’s stop rule stands: do not start P1 without explicit "
        "go-ahead; prefer a new cohort or a sharply constrained secondary question that is not "
        "another open search on these 31 subjects."
    )

    # ============================================================
    # 10. Appendix
    # ============================================================
    r.h("10. Appendix", 1)
    r.h("10.1 Repository layout (README §6)", 2)
    r.p(
        "AXYS ML/ contains data/raw/ (original MAT + survey; do not overwrite) and data/processed/ "
        "(31-female MAT); docs/ (dataset blueprint, Phase 0 write-up, this report); scripts/ "
        "(audit_dataset.py through run_phase6.py, run_p01_deviation.py, run_p02_abnormality.py, "
        "run_p03_shape.py, run_p04_event_phases.py, run_p06_coordination.py, run_power_analysis.py); "
        "src/gait_research/ (catalog.py, matio.py, events.py, normalize.py, labels.py; features/; "
        "aggregation/; statistics/; phenotypes/; within_victim/; trajectories/; similarity/ with "
        "deviation.py, abnormality.py, shape_space.py, event_phases.py, coordination_crp.py, "
        "power_analysis.py, load.py); tests/phase0 … phase6 and tests/similarity/; results/phase0 "
        "… phase6 and results/similarity/ with p01_deviation/, p02_abnormality/, p03_shape/, "
        "p04_event_phases/, p06_coordination/, power_analysis/, and p0_synthesis.md."
    )
    r.p(
        "Do not modify a completed phase’s pipeline or certified outputs when adding later work. "
        "Similarity P0 only reads Phase 1–4 artifacts."
    )
    r.h("10.2 How to run (README §7)", 2)
    r.p(
        "From the project root, with src/ on PYTHONPATH (tests use tests/conftest.py; scripts prepend src/)."
    )
    r.p(
        "Phases 0–6 (frozen; re-run only to regenerate certified outputs): python scripts/audit_dataset.py; "
        "python scripts/extract_gait_cycles.py; python scripts/run_phase2.py; python scripts/certify_phase2.py; "
        "python scripts/run_phase3.py; python scripts/run_phase4.py; python scripts/run_phase5.py; "
        "python scripts/run_phase6.py."
    )
    r.p(
        "Similarity P0 (does not rewrite Phases 0–6): python scripts/run_p01_deviation.py; "
        "python scripts/run_p02_abnormality.py; python scripts/run_p03_shape.py; "
        "python scripts/run_p04_event_phases.py --n-perm 9999; python scripts/run_p06_coordination.py; "
        "python scripts/run_power_analysis.py."
    )
    r.p(
        "Tests: python -m pytest tests/phase0 tests/test_phase1_cycles.py tests/phase2 tests/phase3 "
        "tests/phase4 tests/phase5 tests/phase6 -q; python -m pytest tests/similarity -q."
    )
    r.p(
        "Reproducibility: stochastic steps use seed 20260813 unless a script documents otherwise. "
        "Similarity P0 primary permutations use 9999 subject-label shuffles. The P0.1 MDE simulation "
        "uses 999 perms × 1000 datasets per λ (documented in mde_report.md). Environment: Windows, "
        "Python 3, numpy/pandas/scipy/pyarrow, matplotlib, scikit-learn, tqdm, optional numba. "
        "MATLAB is the original kinematics store; Phase 1+ work from exported npz/parquet."
    )
    r.h("10.3 Where to read results (README §9)", 2)
    r.table(
        ["Phase / P0", "Report", "Certification / lock file"],
        [
            ["0", "results/phase0/audit_report.md", "in that report"],
            ["1", "results/phase1/phase1_report.md", "PASS in report"],
            ["2", "results/phase2/phase2_report.md", "phase2_certification.md"],
            ["3", "results/phase3/phase3_report.md", "phase3_certification.md"],
            ["4", "results/phase4/phase4_report.md", "phase4_certification.md"],
            ["5", "results/phase5/phase5_report.md", "phase5_certification.md"],
            ["6", "results/phase6/phase6_report.md", "phase6_certification.md"],
            ["P0.1", "results/similarity/p01_deviation/p01_report.md", "—"],
            ["P0.2", "results/similarity/p02_abnormality/p02_report.md", "preregistered_features.json"],
            ["P0.3", "results/similarity/p03_shape/p03_report.md", "preregistered_curves.json"],
            ["P0.4", "results/similarity/p04_event_phases/p04_report.md", "preregistered_phases.json"],
            ["P0.6", "results/similarity/p06_coordination/p06_report.md", "preregistered_pairs.json"],
            ["P0.1 MDE", "results/similarity/power_analysis/mde_report.md", "post-hoc; does not alter frozen P0"],
            ["P0 family", "results/similarity/p0_synthesis.md", "decision gate + P1 stop"],
        ],
        widths=[1.2, 3.3, 2.0],
        caption="Table 10.1. Primary report index. Paths are repo-relative.",
    )
    r.p("Useful tables (not reports):")
    r.bullets(
        [
            "Phase 2: subject_features.parquet, feature_catalog.json",
            "Phase 3: candidate_signature.csv, statistics/multiple_testing.csv",
            "Phase 4: phenotype_assignments.csv (no victim column on the assignment table)",
            "Phase 5: similarity/within_victim_similarity.csv",
            "Phase 6: candidate_trajectory_regions.csv (ROBUST / EXPLORATORY / UNSUPPORTED)",
            "P0.1: cosine_matrix.csv, subject_alignment.csv",
            "P0.2: exceedance_matrix.csv, feature_coexceedance_residualized.csv",
            "P0.3: per_curve_similarity_residualized.csv",
            "P0.4: cell_results_residualized.csv, window_multivariate_residualized.csv",
            "P0.6: per_pair_similarity_residualized.csv",
            "MDE: power_analysis/summary.csv",
        ]
    )
    r.h("10.4 Glossary", 2)
    r.table(
        ["Term", "Meaning in this project"],
        [
            ["Subject (n=31)", "Independent sampling unit. Never replace with 880 cycles."],
            ["Cliff’s δ", "Probability that a random victim’s value exceeds a random control’s, minus the reverse; |δ|≥0.33 is the Phase 3 “medium” threshold."],
            ["BH-FDR / q", "Benjamini–Hochberg false-discovery-rate adjusted p; q≤0.10 was the Phase 3 signature and P0 family gate."],
            ["LOSO", "Leave-one-subject-out: drop one person, recompute the statistic or its sign/rank, measure agreement."],
            ["Victim consistency", "Share of victims on the group-difference side of the control median."],
            ["Family-PC", "Within-feature-family PCA, variance-balanced across families, 27-D Phase 4 space."],
            ["Redundancy representative", "One column kept from a Spearman |ρ|≥0.90 cluster (Phase 3: 335 such)."],
            ["Cluster permutation", "Phase 6: permute subject labels; compare observed cluster mass of |t| to the max-mass null."],
            ["CRP", "Continuous relative phase: wrap(φ_proximal−φ_distal); here Hilbert analytic phase of demeaned ax1 angles."],
            ["Jaccard", "Intersection over union of two binary abnormality sets."],
            ["DTW", "Dynamic time warping distance on z-scored curves (P0.3) or unwrapped CRP (P0.6)."],
            ["λ (MDE)", "Shared-offset length ÷ median control residual ||d|| in 27-D; 80% power at λ=0.7285."],
            ["Pseudo-replication", "Treating repeated cycles as independent people; forbidden here."],
            ["ax1/ax2/ax3", "Stored spatial/angle axes; not certified as AP/ML/vertical."],
            ["PiG", "Vicon Plug-in Gait marker/angle model used at capture."],
            ["Seed 20260813", "Default RNG seed for permutations, clustering inits, and MDE."],
        ],
        widths=[1.8, 4.7],
        caption="Table 10.2. Recurring terms.",
    )
    r.h("10.5 Notes on rounding versus primary files", 2)
    r.p(
        "README and p0_synthesis.md sometimes round for readability. This report prefers the "
        "primary CSV/JSON/report table. Documented differences: P0.1 cosine 0.051819… versus "
        "displayed 0.052; P0.1 p 0.7579 versus 0.758; P0.2 Jaccard 0.1906 versus 0.191; P0.3 "
        "Pearson −0.02337 versus −0.023; P0.4 min p 0.0188 versus 0.019; P0.6 circular 0.6494 "
        "versus 0.649 and DTW 17.625 versus 17.63; MDE λ 0.7285 versus 0.73; Phase 4 silhouette "
        "0.41503 versus 0.42; mass 47.7–100.0 kg versus README “~48–100 kg”; height 141.0–167.6 cm "
        "versus “~141–168 cm.” None of these change a decision gate."
    )
    r.p(
        "End of report. Analyses remain frozen. P1 is not started.",
        italic=True,
    )


if __name__ == "__main__":
    build()
